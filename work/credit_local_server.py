from __future__ import annotations

import cgi
import hashlib
import html
import json
import os
import re
import shutil
import sys
import threading
import traceback
import uuid
from difflib import SequenceMatcher
from io import BytesIO
from decimal import Decimal, ROUND_DOWN
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, urlparse

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "work"))
sys.path.insert(0, str(ROOT / "work" / "ocr_deps"))

from ocr_credit_pdf import ocr_images, prepare_and_split, render_pdf
from parse_credit_ocr import parse
from pypdf import PdfReader

UPLOADS = ROOT / "tmp" / "server_uploads"
RESULTS = ROOT / "tmp" / "server_results"
CACHE = ROOT / "tmp" / "ocr_cache"
OCR_CACHE_VERSION = "photo-pages-v3"
APP_VERSION = "多页拍照征信增强版 2026.06.30"
EXPORTS = ROOT / "tmp" / "server_exports"
JOBS: dict[str, dict] = {}
JOBS_LOCK = threading.Lock()


def money(value: float) -> str:
    return f"{float(value or 0):,.0f}"


def money_plain(value: float) -> str:
    amount = Decimal(str(value or 0))
    truncated = amount.quantize(Decimal("0.01"), rounding=ROUND_DOWN)
    if truncated == truncated.to_integral():
        return f"{int(truncated):,}"
    return f"{truncated:,.2f}".rstrip("0").rstrip(".")


def enterprise_yuan(value: float) -> Decimal:
    return Decimal(str(value or 0)) * Decimal("10000")


def enterprise_money(value: float) -> str:
    return money_plain(enterprise_yuan(value))


def enterprise_yuan_number(value: float) -> int | float:
    amount = enterprise_yuan(value)
    if amount == amount.to_integral():
        return int(amount)
    return float(amount)


def set_job(job_id: str, **updates) -> None:
    with JOBS_LOCK:
        job = JOBS.setdefault(job_id, {})
        job.update(updates)


def get_job(job_id: str) -> dict:
    with JOBS_LOCK:
        return dict(JOBS.get(job_id) or {})


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}


def image_series_name(path: Path) -> str:
    """Return a common report name for files such as 张三征信1/张三征信2."""
    stem = path.stem.strip()
    return re.sub(r"(?:[-_\s]*(?:第)?[0-9一二三四五六七八九十]+(?:页)?)$", "", stem).strip("-_ ") or stem


def image_page_order(path: Path) -> tuple[int, str]:
    match = re.search(r"(?:第)?(\d+)(?:页)?$", path.stem)
    return (int(match.group(1)) if match else 10_000, path.name)


def group_uploads(uploads: list[Path]) -> list[list[Path]]:
    """Merge numbered photos from the same report; keep PDFs independent."""
    buckets: dict[str, list[Path]] = {}
    for path in uploads:
        if path.suffix.lower() in IMAGE_SUFFIXES:
            buckets.setdefault(image_series_name(path), []).append(path)
    grouped_names = {name for name, items in buckets.items() if len(items) > 1}
    emitted: set[str] = set()
    groups: list[list[Path]] = []
    for path in uploads:
        if path.suffix.lower() not in IMAGE_SUFFIXES:
            groups.append([path])
            continue
        name = image_series_name(path)
        if name in grouped_names:
            if name not in emitted:
                groups.append(sorted(buckets[name], key=image_page_order))
                emitted.add(name)
        else:
            groups.append([path])
    return groups


def group_label(paths: list[Path]) -> str:
    return image_series_name(paths[0]) if len(paths) > 1 else paths[0].name


def parse_text_result(text_path: Path, uploaded: Path, source_mode: str) -> dict:
    data = parse(text_path)
    data["app_version"] = APP_VERSION
    data["report_html"] = make_report(data)
    data["ocr_text_path"] = str(text_path)
    data["source_mode"] = source_mode
    data["filename"] = uploaded.name
    data["display_name"] = result_name(data, uploaded.name)
    return data


def make_enterprise_report(data: dict) -> str:
    enterprise = data.get("enterprise") or {}
    identity = enterprise.get("identity") or {}
    summary = enterprise.get("summary") or {}
    loans = enterprise.get("loans") or []
    credit_lines = enterprise.get("credit_lines") or []
    off_balance = enterprise.get("off_balance") or []

    attention = float(summary.get("attention_balance") or 0)
    npl = float(summary.get("npl_balance") or 0)
    overdue = float(summary.get("overdue_total") or 0)
    guarantee_balance = float(summary.get("guarantee_balance") or 0)
    risk = "风险偏高" if npl or overdue or attention else ("需关注" if guarantee_balance else "正常关注")

    enterprise_loan_amount = sum(item.get("amount", 0) for item in loans)
    enterprise_loan_balance = sum(item.get("balance", 0) for item in loans)
    enterprise_overdue_total = sum(item.get("overdue_total", 0) for item in loans)
    loan_detail_rows = "".join(
        f"<tr><td>{html.escape(item.get('lender') or '未识别')}</td>"
        f"<td>{html.escape(item.get('type') or '未识别')}</td>"
        f"<td>{html.escape(item.get('guarantee') or '未识别')}</td>"
        f"<td>{enterprise_money(item.get('amount'))}</td><td>{enterprise_money(item.get('balance'))}</td>"
        f"<td>{html.escape(item.get('classification') or '未识别')}</td>"
        f"<td>{enterprise_money(item.get('overdue_total'))}</td><td>{html.escape(str(item.get('overdue_months') or ''))}</td>"
        f"<td>{html.escape(item.get('due_date') or '')}</td></tr>"
        for item in loans
    ) or '<tr><td colspan="9">未识别到未结清借贷明细</td></tr>'
    loan_summary_row = (
        f"<tr class=\"summary-row\"><th colspan=\"3\">合计</th>"
        f"<th>{enterprise_money(enterprise_loan_amount)}</th><th>{enterprise_money(enterprise_loan_balance)}</th>"
        f"<th></th><th>{enterprise_money(enterprise_overdue_total)}</th><th></th><th></th></tr>"
    )
    loan_rows = loan_detail_rows + loan_summary_row
    enterprise_overdue_loans = [
        item for item in loans
        if float(item.get("overdue_total") or 0) > 0
        or str(item.get("overdue_months") or "").strip() not in {"", "0", "-", "--"}
    ]
    enterprise_overdue_rows = "".join(
        f"<tr><td>{html.escape(item.get('lender') or '未识别')}</td>"
        f"<td>{html.escape(item.get('type') or '未识别')}</td>"
        f"<td>具体月份未披露</td><td>{html.escape(str(item.get('overdue_months') or '未识别'))}</td>"
        f"<td>{enterprise_money(item.get('overdue_total'))}</td>"
        f"<td>{html.escape(item.get('classification') or '未识别')}</td></tr>"
        for item in enterprise_overdue_loans
    ) or '<tr><td colspan="6">未识别到逾期贷款</td></tr>'

    credit_rows = "".join(
        f"<tr><td>{html.escape(item.get('lender') or '未识别')}</td>"
        f"<td>{html.escape(item.get('limit_type') or '未识别')}</td>"
        f"<td>{html.escape(item.get('revolving') or '')}</td>"
        f"<td>{enterprise_money(item.get('limit'))}</td><td>{enterprise_money(item.get('used'))}</td>"
        f"<td>{enterprise_money(item.get('credit_limit'))}</td><td>{html.escape(item.get('due_date') or '')}</td></tr>"
        for item in credit_lines
    ) or '<tr><td colspan="7">未识别到授信协议明细</td></tr>'

    off_rows = "".join(
        f"<tr><td>{html.escape(item.get('lender') or '未识别')}</td>"
        f"<td>{html.escape(item.get('type') or '未识别')}</td>"
        f"<td>{html.escape(item.get('classification') or '未识别')}</td>"
        f"<td>{item.get('account_count') or 0}</td><td>{enterprise_money(item.get('balance'))}</td></tr>"
        for item in off_balance
    ) or '<tr><td colspan="5">未识别到表外/担保业务</td></tr>'

    return f"""
      <article class="report-body">
      <div class="report-title">
        <div>
          <h2>企业征信评估报告</h2>
          <p>报告基准日：{html.escape(identity.get("report_date") or "未识别")}；综合判断：<strong>{risk}</strong>。</p>
        </div>
        <span class="pill ok">企业征信模式</span>
      </div>
      <h3>企业基本信息</h3>
      <table><tbody>
      <tr><th>企业名称</th><td>{html.escape(identity.get('company_name') or '未识别')}</td><th>统一社会信用代码</th><td>{html.escape(identity.get('credit_code') or '未识别')}</td></tr>
      <tr><th>中征码</th><td>{html.escape(identity.get('zhongzheng_code') or '未识别')}</td><th>组织机构代码</th><td>{html.escape(identity.get('org_code') or '未识别')}</td></tr>
      <tr><th>法定代表人</th><td>{html.escape(identity.get('legal_representative') or '未识别')}</td><th>实际控制人</th><td>{html.escape(identity.get('actual_controller') or '未识别')}</td></tr>
      <tr><th>注册资本</th><td>{html.escape(identity.get('registered_capital') or '未识别')}</td><th>行业/规模</th><td>{html.escape((identity.get('industry') or '未识别') + ' / ' + (identity.get('enterprise_scale') or '未识别'))}</td></tr>
      </tbody></table>
      <h3>信贷概要</h3>
      <table><thead><tr><th>借贷余额（元）</th><th>担保/表外余额（元）</th><th>关注类余额（元）</th><th>不良类余额（元）</th><th>逾期本金（元）</th><th>逾期利息及其他（元）</th><th>逾期总额（元）</th></tr></thead><tbody>
      <tr><td>{enterprise_money(summary.get('loan_balance'))}</td><td>{enterprise_money(summary.get('guarantee_balance'))}</td><td>{enterprise_money(summary.get('attention_balance'))}</td><td>{enterprise_money(summary.get('npl_balance'))}</td><td>{enterprise_money(summary.get('overdue_principal'))}</td><td>{enterprise_money(summary.get('overdue_interest'))}</td><td>{enterprise_money(summary.get('overdue_total'))}</td></tr>
      </tbody></table>
      <h3>未结清借贷明细</h3>
      <table><thead><tr><th>授信机构</th><th>业务种类</th><th>担保方式</th><th>借款金额（元）</th><th>余额（元）</th><th>五级分类</th><th>逾期总额（元）</th><th>逾期月数</th><th>到期日</th></tr></thead><tbody>{loan_rows}</tbody></table>
      <h3>逾期贷款明细</h3>
      <table><thead><tr><th>授信机构</th><th>业务种类</th><th>逾期时间</th><th>逾期月数</th><th>逾期总额（元）</th><th>五级分类</th></tr></thead><tbody>{enterprise_overdue_rows}</tbody></table>
      <h3>授信信息</h3>
      <table><thead><tr><th>授信机构</th><th>额度类型</th><th>是否循环</th><th>授信额度（元）</th><th>已用额度（元）</th><th>授信限额（元）</th><th>到期日</th></tr></thead><tbody>{credit_rows}</tbody></table>
      <h3>表外及担保业务</h3>
      <table><thead><tr><th>机构</th><th>业务种类</th><th>五级分类</th><th>账户数</th><th>余额（元）</th></tr></thead><tbody>{off_rows}</tbody></table>
      <h3>评估意见</h3>
      <ul>
        <li>企业当前借贷余额约 {enterprise_money(summary.get('loan_balance'))} 元；企业征信原文金额按万元口径识别后已换算为元。</li>
        <li>关注类、逾期或不良金额应作为授信审查重点；本报告识别到逾期总额 {enterprise_money(summary.get('overdue_total'))} 元。</li>
        <li>表外及担保业务余额约 {enterprise_money(summary.get('guarantee_balance'))} 元，应结合银行承兑、保函、保证保险等风险敞口复核。</li>
        <li>扫描版企业征信由 OCR 识别，机构名称、金额和五级分类建议与原始页复核。</li>
      </ul>
      </article>
    """


def personal_overdue_entries(loans: list[dict]) -> list[dict]:
    entries = []
    seen = set()
    for loan in loans:
        events = list(loan.get("overdue_events") or [])
        if not events and (loan.get("current_overdue_periods") or loan.get("current_overdue_amount")):
            events = [{
                "month": "报告基准月",
                "status": f"当前逾期{int(loan.get('current_overdue_periods') or 0)}期",
                "amount": loan.get("current_overdue_amount") or 0,
                "current": True,
            }]
        for event in events:
            entry = {
                "lender": loan.get("lender") or "未识别",
                "type": loan.get("type") or "未识别",
                "month": event.get("month") or "月份未识别",
                "status": event.get("status") or "存在逾期",
                "amount": float(event.get("amount") or 0),
                "current": bool(event.get("current")),
            }
            key = (entry["lender"], entry["type"], entry["month"], entry["status"], entry["amount"], entry["current"])
            if key not in seen:
                seen.add(key)
                entries.append(entry)
    return sorted(entries, key=lambda item: (not item["current"], item["month"], item["lender"]), reverse=False)


def make_report(data: dict) -> str:
    if (data.get("report_type") or {}).get("kind") == "enterprise":
        return make_enterprise_report(data)

    summary = data["summary"]
    card = summary["card_summary"]
    guarantee = summary["guarantee_summary"]
    guarantees = data.get("guarantees") or []
    buckets = data["inquiry_buckets"]
    all_loans = data.get("loans", [])
    other_business = data.get("other_business") or []
    active_loans = [item for item in all_loans if item.get("balance", 0) > 0]
    loan_balance = sum(item.get("balance", 0) for item in active_loans)
    monthly = sum(item.get("monthly_payment", 0) for item in active_loans)
    card_rate = card["used"] / card["total_limit"] if card["total_limit"] else 0
    card_last6 = money(card["last6_avg_used"]) if card.get("last6_avg_used") else "简版未披露"

    personal_limit = sum(item.get("limit", 0) for item in all_loans)
    loan_detail_rows = "".join(
        f"<tr><td>{html.escape(item['lender'])}</td><td>{html.escape(item['type'])}</td>"
        f"<td>{html.escape(item['guarantee'])}</td><td>{money(item['limit'])}</td>"
        f"<td>{money(item['balance'])}</td><td>{html.escape(str(item['remaining_terms']))}</td>"
        f"<td>{money(item['monthly_payment'])}</td></tr>"
        for item in all_loans
    ) or '<tr><td colspan="7">未识别到贷款账户</td></tr>'
    loan_summary_row = (
        f"<tr class=\"summary-row\"><th colspan=\"3\">合计</th>"
        f"<th>{money(personal_limit)}</th><th>{money(loan_balance)}</th><th></th><th>{money(monthly)}</th></tr>"
    )
    loan_rows = loan_detail_rows + loan_summary_row
    other_business_rows = "".join(
        f"<tr><td>{html.escape(item.get('lender') or '未识别')}</td>"
        f"<td>{html.escape(item.get('type') or '未识别')}</td>"
        f"<td>{money(item.get('amount'))}</td><td>{money(item.get('balance'))}</td>"
        f"<td>{html.escape(item.get('due_date') or '未披露')}</td></tr>"
        for item in other_business
    ) or '<tr><td colspan="5">未识别到其他信贷业务</td></tr>'
    overdue_entries = personal_overdue_entries(all_loans)
    overdue_rows = "".join(
        f"<tr><td>{html.escape(item['lender'])}</td><td>{html.escape(item['type'])}</td>"
        f"<td>{html.escape(item['month'])}</td><td>{html.escape(item['status'])}</td>"
        f"<td>{money(item['amount']) if item['amount'] else '未披露'}</td>"
        f"<td>{'当前逾期' if item['current'] else '历史逾期'}</td></tr>"
        for item in overdue_entries
    ) or '<tr><td colspan="6">未识别到逾期贷款</td></tr>'

    card_rows = "".join(
        f"<tr><td>{html.escape(item['lender'])}</td><td>{money(item['limit'])}</td><td>{money(item['used'])}</td>"
        f"<td>{(item['used'] / item['limit'] if item['limit'] else 0):.1%}</td></tr>"
        for item in data.get("cards", [])
    )
    card_detail = (
        f"<h3>信用卡明细</h3><table><thead><tr><th>发卡机构</th><th>信用额度</th><th>已使用额度</th><th>使用率</th></tr></thead><tbody>{card_rows}</tbody></table>"
        if card_rows else ""
    )

    risk = "风险偏高" if overdue_entries or card_rate >= 0.85 or guarantee["balance"] > 0 else "需关注"
    guarantee_rows = "".join(
        f"<tr><td>{html.escape(item.get('main_borrower') or '未识别')}</td>"
        f"<td>{html.escape(item.get('lender') or '未识别')}</td>"
        f"<td>{money(item.get('guarantee_amount'))}</td><td>{money(item.get('balance'))}</td></tr>"
        for item in guarantees
    )
    if not guarantee_rows and (guarantee["balance"] or guarantee["guarantee_amount"]):
        guarantee_rows = (
            f"<tr><td>{html.escape(guarantee.get('main_borrower') or '未识别，需复核原始报告')}</td>"
            f"<td>未识别</td><td>{money(guarantee['guarantee_amount'])}</td><td>{money(guarantee['balance'])}</td></tr>"
        )
    if not guarantee_rows:
        guarantee_rows = '<tr><td colspan="4">未识别到对外担保责任</td></tr>'
    guarantee_note = html.escape(guarantee.get("main_borrower") or "未识别")
    guarantee_tip = (
        f"主业务借款人：{guarantee_note}；对外担保余额 {money(guarantee['balance'])} 元，应作为重要或有负债纳入审查。"
        if guarantee["balance"] or guarantee["guarantee_amount"]
        else "未识别到对外担保余额。"
    )
    return f"""
      <article class="report-body">
      <div class="report-title">
        <div>
          <h2>征信评估报告</h2>
          <p>报告基准日：{html.escape(summary.get("report_date") or "未识别")}；综合判断：<strong>{risk}</strong>。</p>
        </div>
      </div>
      <h3>贷款明细</h3>
      <table><thead><tr><th>贷款机构名称</th><th>业务种类</th><th>担保方式</th><th>账户授信额度</th><th>贷款余额</th><th>剩余还款期数</th><th>本月应还款</th></tr></thead><tbody>{loan_rows}</tbody></table>
      <p>有余额贷款合计：{money(loan_balance)} 元；本月应还款合计约：{money(monthly)} 元。</p>
      <h3>其他信贷业务</h3>
      <table><thead><tr><th>业务机构</th><th>业务类型</th><th>业务金额</th><th>当前余额</th><th>到期日</th></tr></thead><tbody>{other_business_rows}</tbody></table>
      <h3>逾期贷款明细</h3>
      <table><thead><tr><th>贷款机构名称</th><th>业务种类</th><th>逾期月份</th><th>逾期程度</th><th>逾期金额</th><th>状态</th></tr></thead><tbody>{overdue_rows}</tbody></table>
      <h3>信用卡汇总</h3>
      <table><thead><tr><th>授信总额</th><th>已使用额度</th><th>近6个月平均使用额度</th><th>使用率</th></tr></thead><tbody>
      <tr><td>{money(card['total_limit'])}</td><td>{money(card['used'])}</td><td>{card_last6}</td><td>{card_rate:.1%}</td></tr>
      </tbody></table>
      {card_detail}
      <h3>对外担保</h3>
      <table><thead><tr><th>主业务借款人</th><th>贷款机构</th><th>担保/责任金额</th><th>贷款余额</th></tr></thead><tbody>
      {guarantee_rows}
      </tbody></table>
      <h3>审批查询统计</h3>
      <table><thead><tr><th>期间</th><th>贷款审批</th><th>信用卡审批</th></tr></thead><tbody>
      <tr><td>近一年</td><td>{buckets['year']['loan']}</td><td>{buckets['year']['card']}</td></tr>
      <tr><td>近半年</td><td>{buckets['half_year']['loan']}</td><td>{buckets['half_year']['card']}</td></tr>
      <tr><td>近三个月</td><td>{buckets['quarter']['loan']}</td><td>{buckets['quarter']['card']}</td></tr>
      <tr><td>近一个月</td><td>{buckets['month']['loan']}</td><td>{buckets['month']['card']}</td></tr>
      </tbody></table>
      <h3>评估意见</h3>
      <ul>
        <li>信用卡使用率约 {card_rate:.1%}，若接近或超过 85%，应重点核实资金周转压力。</li>
        <li>{f'识别到 {len(overdue_entries)} 条贷款逾期记录，应重点复核逾期月份、持续时间及是否已结清。' if overdue_entries else '未识别到贷款逾期记录。'}</li>
        <li>{guarantee_tip}</li>
        <li>当前报告由本地 OCR 识别扫描件后生成，关键金额、机构名称和主业务借款人仍建议与原始页复核。</li>
      </ul>
      </article>
    """


def extract_pdf_text(pdf_path: Path) -> str:
    try:
        reader = PdfReader(str(pdf_path))
        return "\n\n".join(
            f"【第 {i + 1} 页】\n{page.extract_text() or ''}"
            for i, page in enumerate(reader.pages)
        )
    except Exception:
        return ""


def result_name(data: dict, fallback: str) -> str:
    if (data.get("report_type") or {}).get("kind") == "enterprise":
        return ((data.get("enterprise") or {}).get("identity") or {}).get("company_name") or Path(fallback).stem
    stem = image_series_name(Path(fallback))
    filename_name = re.sub(r"(?:个人)?征信(?:报告)?$", "", stem).strip("-_ ") or stem
    recognized = (data.get("identity") or {}).get("name") or ""
    if recognized and (recognized in filename_name or SequenceMatcher(None, recognized, filename_name).ratio() >= 0.6):
        return recognized
    return filename_name


def result_summary(data: dict) -> dict:
    kind = (data.get("report_type") or {}).get("kind")
    if kind == "enterprise":
        ent = data.get("enterprise") or {}
        identity = ent.get("identity") or {}
        summary = ent.get("summary") or {}
        credit_lines = ent.get("credit_lines") or []
        overdue = float(summary.get("overdue_total") or 0)
        attention = float(summary.get("attention_balance") or 0)
        npl = float(summary.get("npl_balance") or 0)
        risk = "风险偏高" if overdue or attention or npl else "正常关注"
        return {
            "name": identity.get("company_name") or data.get("display_name") or data.get("filename") or "未识别",
            "type": "企业征信",
            "report_date": identity.get("report_date") or "",
            "risk": risk,
            "personal_loan_balance": 0,
            "card_used": 0,
            "card_limit": 0,
            "personal_guarantee_balance": 0,
            "enterprise_loan_balance": enterprise_yuan_number(summary.get("loan_balance", 0)),
            "enterprise_credit_used": enterprise_yuan_number(sum(item.get("used", 0) for item in credit_lines)),
            "enterprise_guarantee_balance": enterprise_yuan_number(summary.get("guarantee_balance", 0)),
            "overdue_total": enterprise_yuan_number(overdue),
            "loan_count": len(ent.get("loans") or []),
            "source_mode": data.get("source_mode") or "OCR识别",
        }

    summary = data.get("summary") or {}
    card = summary.get("card_summary") or {}
    guarantee = summary.get("guarantee_summary") or {}
    active_loans = [item for item in data.get("loans", []) if item.get("balance", 0) > 0]
    loan_balance = sum(item.get("balance", 0) for item in active_loans)
    card_rate = (card.get("used", 0) / card.get("total_limit", 0)) if card.get("total_limit") else 0
    overdue_entries = personal_overdue_entries(data.get("loans") or [])
    risk = "风险偏高" if overdue_entries or card_rate >= 0.85 or guarantee.get("balance", 0) else "需关注"
    return {
        "name": data.get("display_name") or data.get("filename") or "未识别",
        "type": "个人征信",
        "report_date": summary.get("report_date") or "",
        "risk": risk,
        "personal_loan_balance": loan_balance,
        "card_used": card.get("used", 0),
        "card_limit": card.get("total_limit", 0),
        "personal_guarantee_balance": guarantee.get("balance", 0),
        "enterprise_loan_balance": 0,
        "enterprise_credit_used": 0,
        "enterprise_guarantee_balance": 0,
        "overdue_total": 0,
        "loan_count": len(data.get("loans") or []),
        "source_mode": data.get("source_mode") or "OCR识别",
    }


def make_batch_report(results: list[dict]) -> str:
    summaries = [result_summary(item) for item in results]
    personal_count = sum(1 for item in summaries if item["type"] == "个人征信")
    enterprise_count = sum(1 for item in summaries if item["type"] == "企业征信")
    high_risk = sum(1 for item in summaries if item["risk"] == "风险偏高")
    personal_loan_total = sum(item["personal_loan_balance"] for item in summaries)
    card_used_total = sum(item["card_used"] for item in summaries)
    personal_guarantee_total = sum(item["personal_guarantee_balance"] for item in summaries)
    enterprise_loan_total = sum(item["enterprise_loan_balance"] for item in summaries)
    enterprise_credit_used_total = sum(item["enterprise_credit_used"] for item in summaries)
    enterprise_guarantee_total = sum(item["enterprise_guarantee_balance"] for item in summaries)
    overdue_total = sum(item["overdue_total"] for item in summaries)

    rows = "".join(
        f"<tr><td>{html.escape(item['name'])}</td><td>{item['type']}</td><td>{html.escape(item['report_date'] or '未识别')}</td>"
        f"<td>{item['risk']}</td><td>{item['loan_count']}</td><td>{money(item['personal_loan_balance'])}</td>"
        f"<td>{money(item['card_used'])}/{money(item['card_limit'])}</td><td>{money(item['personal_guarantee_balance'])}</td>"
        f"<td>{money(item['enterprise_loan_balance'])}</td><td>{money(item['enterprise_credit_used'])}</td>"
        f"<td>{money(item['enterprise_guarantee_balance'])}</td><td>{money(item['overdue_total'])}</td>"
        f"<td>{html.escape(item['source_mode'])}</td></tr>"
        for item in summaries
    )

    single_reports = []
    for index, item in enumerate(results, start=1):
        summary = result_summary(item)
        label = (item.get("report_type") or {}).get("label") or "征信报告"
        source = item.get("source_mode") or "OCR识别"
        single_reports.append(
            f"""
            <section class="single-report">
              <div class="single-head">
                <div>
                  <h3>第 {index} 份报告：{html.escape(summary['name'])}</h3>
                  <p>{html.escape(label)}；来源：{html.escape(source)}</p>
                </div>
                <span class="pill">{html.escape(summary['risk'])}</span>
              </div>
              {item.get('report_html', '')}
            </section>
            """
        )
    details = "".join(single_reports)

    return f"""
      <article class="report-body">
      <div class="report-title">
        <div>
          <h2>批量征信综合分析报告</h2>
          <p>共分析 <strong>{len(results)}</strong> 份征信：个人 {personal_count} 份，企业 {enterprise_count} 份；风险偏高 {high_risk} 份。</p>
        </div>
        <span class="pill ok">批量汇总</span>
      </div>
      <h3>综合汇总</h3>
      <table><thead><tr><th>个人贷款余额合计（元）</th><th>信用卡已用合计（元）</th><th>个人担保余额合计（元）</th><th>企业借贷余额合计（元）</th><th>企业授信已用合计（元）</th><th>企业表外/担保余额合计（元）</th><th>企业逾期总额合计（元）</th></tr></thead><tbody>
      <tr><td>{money(personal_loan_total)}</td><td>{money(card_used_total)}</td><td>{money(personal_guarantee_total)}</td><td>{money(enterprise_loan_total)}</td><td>{money(enterprise_credit_used_total)}</td><td>{money(enterprise_guarantee_total)}</td><td>{money(overdue_total)}</td></tr>
      </tbody></table>
      <h3>文件汇总表</h3>
      <table><thead><tr><th>客户/企业名称</th><th>类型</th><th>报告日期</th><th>风险</th><th>未结清笔数</th><th>个人贷款余额（元）</th><th>信用卡已用/授信（元）</th><th>个人担保余额（元）</th><th>企业借贷余额（元）</th><th>企业授信已用（元）</th><th>企业表外/担保余额（元）</th><th>企业逾期总额（元）</th><th>来源</th></tr></thead><tbody>{rows}</tbody></table>
      <h3>综合评估意见</h3>
      <ul>
        <li>个人征信和企业征信均按“元”展示；企业征信原文金额按万元口径识别后已换算为元。</li>
        <li>建议优先复核风险偏高、信用卡使用率较高、存在个人担保、企业关注/逾期/表外敞口的主体。</li>
        <li>扫描版文件由 OCR 识别，关键机构名称、金额和五级分类仍建议结合原始征信复核。</li>
      </ul>
      <h3>每份征信单独报告</h3>
      {details}
      </article>
    """


def make_batch_response(results: list[dict]) -> dict:
    summaries = [result_summary(item) for item in results]
    return {
        "batch": True,
        "app_version": APP_VERSION,
        "report_type": {"kind": "batch", "label": "批量征信"},
        "results": results,
        "summaries": summaries,
        "text_length": sum(item.get("text_length", 0) for item in results),
        "report_html": make_batch_report(results),
        "aggregate": {
            "file_count": len(results),
            "personal_count": sum(1 for item in summaries if item["type"] == "个人征信"),
            "enterprise_count": sum(1 for item in summaries if item["type"] == "企业征信"),
            "high_risk_count": sum(1 for item in summaries if item["risk"] == "风险偏高"),
            "personal_loan_total": sum(item["personal_loan_balance"] for item in summaries),
            "enterprise_loan_total": sum(item["enterprise_loan_balance"] for item in summaries),
            "enterprise_guarantee_total": sum(item["enterprise_guarantee_balance"] for item in summaries),
        },
    }


def safe_filename(name: str) -> str:
    text = re.sub(r'[\\/:*?"<>|\s]+', "_", name or "征信评估报告").strip("_")
    return text[:80] or "征信评估报告"


def export_title(data: dict) -> str:
    if data.get("batch"):
        return "批量征信综合分析报告"
    kind = (data.get("report_type") or {}).get("kind")
    if kind == "enterprise":
        name = ((data.get("enterprise") or {}).get("identity") or {}).get("company_name") or data.get("display_name") or ""
        return f"{name}企业征信评估报告" if name else "企业征信评估报告"
    name = data.get("display_name") or data.get("filename") or ""
    return f"{Path(name).stem}征信评估报告" if name else "征信评估报告"


def to_export_number(value, enterprise_unit: bool = False):
    if enterprise_unit:
        return enterprise_yuan_number(value or 0)
    try:
        return float(value or 0)
    except Exception:
        return 0


def table_block(title: str, headers: list[str], rows: list[list]) -> dict:
    return {"title": title, "headers": headers, "rows": rows or [["无数据"] + [""] * (len(headers) - 1)]}


def personal_export_blocks(data: dict) -> list[dict]:
    summary = data.get("summary") or {}
    card = summary.get("card_summary") or {}
    guarantee = summary.get("guarantee_summary") or {}
    buckets = data.get("inquiry_buckets") or {}
    loans = data.get("loans") or []
    cards = data.get("cards") or []
    guarantees = data.get("guarantees") or []
    active_loans = [item for item in loans if float(item.get("balance") or 0) > 0]
    overdue_entries = personal_overdue_entries(loans)
    blocks = [
        table_block("基本概要", ["项目", "数值"], [
            ["报告类型", "个人征信"],
            ["报告日期", summary.get("report_date") or "未识别"],
            ["贷款余额合计", sum(float(item.get("balance") or 0) for item in active_loans)],
            ["信用卡授信总额", card.get("total_limit") or 0],
            ["信用卡已使用额度", card.get("used") or 0],
            ["对外担保余额", guarantee.get("balance") or 0],
            ["来源", data.get("source_mode") or "OCR识别"],
        ]),
        table_block("贷款明细", ["贷款机构名称", "业务种类", "担保方式", "账户授信额度", "贷款余额", "剩余还款期数", "本月应还款"], [
            [
                item.get("lender") or "未识别",
                item.get("type") or "未识别",
                item.get("guarantee") or "未识别",
                to_export_number(item.get("limit")),
                to_export_number(item.get("balance")),
                item.get("remaining_terms") or "",
                to_export_number(item.get("monthly_payment")),
            ]
            for item in loans
        ]),
        table_block("逾期贷款明细", ["贷款机构名称", "业务种类", "逾期月份", "逾期程度", "逾期金额", "状态"], [
            [
                item.get("lender") or "未识别",
                item.get("type") or "未识别",
                item.get("month") or "月份未识别",
                item.get("status") or "存在逾期",
                to_export_number(item.get("amount")) if item.get("amount") else "未披露",
                "当前逾期" if item.get("current") else "历史逾期",
            ]
            for item in overdue_entries
        ]),
        table_block("信用卡明细", ["发卡机构", "信用额度", "已使用额度", "使用率"], [
            [
                item.get("lender") or "未识别",
                to_export_number(item.get("limit")),
                to_export_number(item.get("used")),
                f"{(float(item.get('used') or 0) / float(item.get('limit') or 1)):.1%}" if float(item.get("limit") or 0) else "0.0%",
            ]
            for item in cards
        ]),
        table_block("对外担保", ["主业务借款人", "贷款机构", "担保/责任金额", "贷款余额"], [
            [
                item.get("main_borrower") or "未识别",
                item.get("lender") or "未识别",
                to_export_number(item.get("guarantee_amount")),
                to_export_number(item.get("balance")),
            ]
            for item in guarantees
        ]),
        table_block("审批查询统计", ["期间", "贷款审批", "信用卡审批"], [
            ["近一年", (buckets.get("year") or {}).get("loan", 0), (buckets.get("year") or {}).get("card", 0)],
            ["近半年", (buckets.get("half_year") or {}).get("loan", 0), (buckets.get("half_year") or {}).get("card", 0)],
            ["近三个月", (buckets.get("quarter") or {}).get("loan", 0), (buckets.get("quarter") or {}).get("card", 0)],
            ["近一个月", (buckets.get("month") or {}).get("loan", 0), (buckets.get("month") or {}).get("card", 0)],
        ]),
    ]
    return blocks


def enterprise_export_blocks(data: dict) -> list[dict]:
    ent = data.get("enterprise") or {}
    identity = ent.get("identity") or {}
    summary = ent.get("summary") or {}
    loans = ent.get("loans") or []
    credit_lines = ent.get("credit_lines") or []
    off_balance = ent.get("off_balance") or []
    return [
        table_block("企业基本信息", ["项目", "数值"], [
            ["企业名称", identity.get("company_name") or "未识别"],
            ["统一社会信用代码", identity.get("credit_code") or "未识别"],
            ["中征码", identity.get("zhongzheng_code") or "未识别"],
            ["报告日期", identity.get("report_date") or "未识别"],
            ["法定代表人", identity.get("legal_representative") or "未识别"],
            ["实际控制人", identity.get("actual_controller") or "未识别"],
            ["来源", data.get("source_mode") or "OCR识别"],
        ]),
        table_block("企业信贷概要", ["项目", "金额"], [
            ["借贷余额", to_export_number(summary.get("loan_balance"), True)],
            ["担保/表外余额", to_export_number(summary.get("guarantee_balance"), True)],
            ["关注类余额", to_export_number(summary.get("attention_balance"), True)],
            ["不良类余额", to_export_number(summary.get("npl_balance"), True)],
            ["逾期总额", to_export_number(summary.get("overdue_total"), True)],
        ]),
        table_block("未结清借贷明细", ["授信机构", "业务种类", "担保方式", "借款金额", "余额", "五级分类", "逾期总额", "逾期月数", "到期日"], [
            [
                item.get("lender") or "未识别",
                item.get("type") or "未识别",
                item.get("guarantee") or "未识别",
                to_export_number(item.get("amount"), True),
                to_export_number(item.get("balance"), True),
                item.get("classification") or "未识别",
                to_export_number(item.get("overdue_total"), True),
                item.get("overdue_months") or "",
                item.get("due_date") or "",
            ]
            for item in loans
        ]),
        table_block("逾期贷款明细", ["授信机构", "业务种类", "逾期时间", "逾期月数", "逾期总额", "五级分类"], [
            [
                item.get("lender") or "未识别",
                item.get("type") or "未识别",
                "具体月份未披露",
                item.get("overdue_months") or "未识别",
                to_export_number(item.get("overdue_total"), True),
                item.get("classification") or "未识别",
            ]
            for item in loans
            if float(item.get("overdue_total") or 0) > 0
            or str(item.get("overdue_months") or "").strip() not in {"", "0", "-", "--"}
        ]),
        table_block("授信信息", ["授信机构", "额度类型", "是否循环", "授信额度", "已用额度", "授信限额", "到期日"], [
            [
                item.get("lender") or "未识别",
                item.get("limit_type") or "未识别",
                item.get("revolving") or "",
                to_export_number(item.get("limit"), True),
                to_export_number(item.get("used"), True),
                to_export_number(item.get("credit_limit"), True),
                item.get("due_date") or "",
            ]
            for item in credit_lines
        ]),
        table_block("表外及担保业务", ["机构", "业务种类", "五级分类", "账户数", "余额"], [
            [
                item.get("lender") or "未识别",
                item.get("type") or "未识别",
                item.get("classification") or "未识别",
                item.get("account_count") or 0,
                to_export_number(item.get("balance"), True),
            ]
            for item in off_balance
        ]),
    ]


def result_export_blocks(data: dict) -> list[dict]:
    if (data.get("report_type") or {}).get("kind") == "enterprise":
        return enterprise_export_blocks(data)
    return personal_export_blocks(data)


def export_blocks(data: dict) -> list[dict]:
    if not data.get("batch"):
        return result_export_blocks(data)
    summaries = data.get("summaries") or [result_summary(item) for item in data.get("results", [])]
    blocks = [
        table_block("批量汇总", ["项目", "数值"], [
            ["文件数量", len(data.get("results") or [])],
            ["个人征信份数", (data.get("aggregate") or {}).get("personal_count", 0)],
            ["企业征信份数", (data.get("aggregate") or {}).get("enterprise_count", 0)],
            ["风险偏高份数", (data.get("aggregate") or {}).get("high_risk_count", 0)],
        ]),
        table_block("文件汇总表", ["客户/企业名称", "类型", "报告日期", "风险", "未结清笔数", "个人贷款余额", "信用卡已用", "信用卡授信", "个人担保余额", "企业借贷余额", "企业授信已用", "企业表外/担保余额", "企业逾期总额", "来源"], [
            [
                item.get("name") or "未识别",
                item.get("type") or "",
                item.get("report_date") or "",
                item.get("risk") or "",
                item.get("loan_count") or 0,
                item.get("personal_loan_balance") or 0,
                item.get("card_used") or 0,
                item.get("card_limit") or 0,
                item.get("personal_guarantee_balance") or 0,
                item.get("enterprise_loan_balance") or 0,
                item.get("enterprise_credit_used") or 0,
                item.get("enterprise_guarantee_balance") or 0,
                item.get("overdue_total") or 0,
                item.get("source_mode") or "",
            ]
            for item in summaries
        ]),
    ]
    for idx, item in enumerate(data.get("results") or [], start=1):
        name = result_summary(item).get("name") or f"第{idx}份"
        for block in result_export_blocks(item):
            blocks.append({**block, "title": f"第 {idx} 份 - {name} - {block['title']}"})
    return blocks


def set_docx_font(run, size: int | None = None, bold: bool = False):
    run.font.name = "Arial"
    if size:
        from docx.shared import Pt
        run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")


def build_docx_export(data: dict) -> bytes:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)
    style._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_docx_font(title.add_run(export_title(data)), 18, True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_docx_font(sub.add_run("由征信分析平台生成，关键金额与机构名称建议结合原件复核。"), 9)
    for block in export_blocks(data):
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
        set_docx_font(heading.add_run(block["title"]), 12, True)
        table = doc.add_table(rows=1, cols=len(block["headers"]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for cell, header in zip(table.rows[0].cells, block["headers"]):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_docx_font(cell.paragraphs[0].add_run(str(header)), 8, True)
        for row in block["rows"]:
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                para = cell.paragraphs[0]
                text = money_plain(value) if isinstance(value, (int, float, Decimal)) and not isinstance(value, bool) else str(value or "")
                set_docx_font(para.add_run(text), 8)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_xlsx_export(data: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    default = wb.active
    wb.remove(default)
    header_fill = PatternFill("solid", fgColor="EAF1F8")
    header_font = Font(bold=True, color="1F334A")

    def unique_sheet_name(name: str) -> str:
        base = re.sub(r"[\[\]:*?/\\]", "", name)[:28] or "报告"
        candidate = base
        idx = 1
        while candidate in wb.sheetnames:
            suffix = f"_{idx}"
            candidate = f"{base[:31-len(suffix)]}{suffix}"
            idx += 1
        return candidate

    for block in export_blocks(data):
        ws = wb.create_sheet(unique_sheet_name(block["title"]))
        ws.append(block["headers"])
        for row in block["rows"]:
            ws.append(row)
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                if isinstance(cell.value, (int, float)):
                    cell.number_format = "#,##0.00" if abs(float(cell.value) - int(float(cell.value))) > 0 else "#,##0"
        for col in range(1, ws.max_column + 1):
            letter = get_column_letter(col)
            max_len = max(len(str(ws.cell(row=row, column=col).value or "")) for row in range(1, min(ws.max_row, 80) + 1))
            ws.column_dimensions[letter].width = min(max(max_len + 3, 12), 34)
    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()


def build_pdf_export(data: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = "Helvetica"
    for font_path in [
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ]:
        try:
            pdfmetrics.registerFont(TTFont("CreditExportCN", font_path))
            font_name = "CreditExportCN"
            break
        except Exception:
            continue
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CnTitle", parent=styles["Title"], fontName=font_name, fontSize=18, leading=24, alignment=1)
    h_style = ParagraphStyle("CnHeading", parent=styles["Heading2"], fontName=font_name, fontSize=12, leading=16, spaceBefore=8, spaceAfter=4)
    cell_style = ParagraphStyle("CnCell", parent=styles["BodyText"], fontName=font_name, fontSize=7, leading=9)
    note_style = ParagraphStyle("CnNote", parent=styles["BodyText"], fontName=font_name, fontSize=8, leading=11, alignment=1)
    bio = BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=landscape(A4), leftMargin=10 * mm, rightMargin=10 * mm, topMargin=10 * mm, bottomMargin=10 * mm)
    story = [Paragraph(export_title(data), title_style), Paragraph("由征信分析平台生成，关键金额与机构名称建议结合原件复核。", note_style), Spacer(1, 5 * mm)]
    usable_width = landscape(A4)[0] - 20 * mm
    for block in export_blocks(data):
        story.append(Paragraph(block["title"], h_style))
        table_data = [[Paragraph(str(h), cell_style) for h in block["headers"]]]
        for row in block["rows"]:
            table_data.append([
                Paragraph(money_plain(value) if isinstance(value, (int, float, Decimal)) and not isinstance(value, bool) else str(value or ""), cell_style)
                for value in row
            ])
        col_width = usable_width / max(1, len(block["headers"]))
        table = Table(table_data, colWidths=[col_width] * len(block["headers"]), repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF1F8")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1F334A")),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.extend([table, Spacer(1, 4 * mm)])
    doc.build(story)
    return bio.getvalue()


def build_export(data: dict, fmt: str) -> tuple[bytes, str, str]:
    fmt = (fmt or "").lower()
    title = safe_filename(export_title(data))
    if fmt == "pdf":
        return build_pdf_export(data), f"{title}.pdf", "application/pdf"
    if fmt in {"word", "docx"}:
        return build_docx_export(data), f"{title}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt in {"excel", "xlsx"}:
        return build_xlsx_export(data), f"{title}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    raise ValueError("下载格式不支持。")


def analyze_image_group(images: list[Path], progress=None) -> dict:
    """Analyze numbered photos as pages of one credit report."""
    images = sorted(images, key=image_page_order)
    report_name = image_series_name(images[0])
    job = RESULTS / report_name
    if job.exists():
        shutil.rmtree(job)
    job.mkdir(parents=True, exist_ok=True)
    CACHE.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha256()
    for image in images:
        digest.update(file_hash(image).encode("ascii"))
    cache_text = CACHE / f"{digest.hexdigest()}-{OCR_CACHE_VERSION}.txt"
    text_path = job / "ocr.txt"
    if cache_text.exists():
        if progress:
            progress(82, "复用多页照片识别结果")
        shutil.copyfile(cache_text, text_path)
    else:
        split_images = []
        total = len(images)
        for page_no, image in enumerate(images, start=1):
            if progress:
                progress(8 + int(30 * (page_no - 1) / total), f"增强第 {page_no}/{total} 页照片")
            split_images.extend(prepare_and_split(image, job / "split", page_no, include_wide_splits=False))
        def ocr_progress(done, count, _img):
            if progress:
                progress(40 + int(45 * done / max(count, 1)), f"OCR 识别中 {done}/{count} 页")
        pages = ocr_images(split_images, progress_cb=ocr_progress)
        text = "\n\n".join(f"【第 {page['page']} 页】\n{page['text']}" for page in pages)
        text_path.write_text(text, encoding="utf-8")
        cache_text.write_text(text, encoding="utf-8")
    if progress:
        progress(92, "合并多页并解析字段")
    data = parse_text_result(text_path, images[0], "多页照片增强OCR")
    data["filename"] = "、".join(image.name for image in images)
    data["display_name"] = result_name(data, f"{report_name}.jpg")
    data["source_files"] = [image.name for image in images]
    data["page_count"] = len(images)
    if progress:
        progress(100, "生成合并报告完成")
    return data


def run_analysis_job(job_id: str, uploads: list[Path]) -> None:
    try:
        results = []
        groups = group_uploads(uploads)
        total = len(groups) or 1
        for idx, paths in enumerate(groups, start=1):
            uploaded = paths[0]
            label = group_label(paths)
            base = int((idx - 1) * 100 / total)
            span = 100 / total

            def progress(file_pct, message, idx=idx, label=label):
                overall = min(99, int(base + span * file_pct / 100))
                set_job(
                    job_id,
                    status="running",
                    percent=overall,
                    current_file=idx,
                    total_files=total,
                    filename=label,
                    message=f"第 {idx}/{total} 份：{label} - {message}",
                )

            progress(0, "开始处理")
            if len(paths) > 1:
                results.append(analyze_image_group(paths, progress=progress))
            else:
                results.append(analyze_file(uploaded, progress=progress))

        result = make_batch_response(results) if len(results) > 1 else results[0]
        set_job(job_id, status="done", percent=100, message="分析完成", result=result)
    except Exception as exc:
        set_job(
            job_id,
            status="error",
            percent=100,
            message=str(exc),
            error=str(exc),
            trace=traceback.format_exc(),
        )


def analyze_file(uploaded: Path, progress=None) -> dict:
    job = RESULTS / uploaded.stem
    if job.exists():
        shutil.rmtree(job)
    job.mkdir(parents=True, exist_ok=True)
    suffix = uploaded.suffix.lower()
    photo_spread = False
    CACHE.mkdir(parents=True, exist_ok=True)
    cache_text = CACHE / f"{file_hash(uploaded)}-{OCR_CACHE_VERSION}.txt"

    if progress:
        progress(0, "准备读取文件")
    if suffix == ".pdf":
        if progress:
            progress(5, "检查 PDF 是否有文字层")
        extracted = extract_pdf_text(uploaded)
        text_path = job / "text_layer.txt"
        if len(extracted.strip()) >= 800:
            if progress:
                progress(70, "读取 PDF 文字层")
            text_path.write_text(extracted, encoding="utf-8")
            data = parse_text_result(text_path, uploaded, "PDF文字层")
            if progress:
                progress(100, "生成报告完成")
            return data

        if cache_text.exists():
            if progress:
                progress(80, "复用上次 OCR 识别结果")
            text_path = job / "ocr.txt"
            shutil.copyfile(cache_text, text_path)
            data = parse_text_result(text_path, uploaded, "OCR缓存")
            if progress:
                progress(100, "生成报告完成")
            return data

        split_images = []
        if progress:
            progress(12, "渲染 PDF 页面")
        def render_progress(done, total, message):
            if progress:
                pct = 12 + int(18 * done / max(total, 1))
                progress(pct, message)

        rendered = render_pdf(uploaded, job / "rendered", progress_cb=render_progress)
        if not rendered:
            raise ValueError("PDF 页面渲染失败，无法进入 OCR 识别。")
        if progress:
            progress(24, f"图像增强与拆页，共 {len(rendered)} 页")
        for page_no, image in enumerate(rendered, start=1):
            if progress:
                progress(30 + int(18 * (page_no - 1) / max(len(rendered), 1)), f"图像增强与拆页 {page_no}/{len(rendered)} 页")
            page_images = prepare_and_split(image, job / "split", page_no)
            photo_spread = photo_spread or any("_photo_" in item.name for item in page_images)
            split_images.extend(page_images)
            if progress:
                progress(30 + int(18 * page_no / max(len(rendered), 1)), f"已完成图像增强 {page_no}/{len(rendered)} 页")
        if progress:
            progress(48, f"拆页完成：{len(rendered)} 个PDF页，共 {len(split_images)} 个实际页面")
    elif suffix in IMAGE_SUFFIXES:
        text_path = job / "ocr.txt"
        if cache_text.exists():
            if progress:
                progress(80, "复用上次 OCR 识别结果")
            shutil.copyfile(cache_text, text_path)
            data = parse_text_result(text_path, uploaded, "OCR缓存")
            if progress:
                progress(100, "生成报告完成")
            return data

        split_images = []
        if progress:
            progress(15, "图像增强")
        split_images.extend(prepare_and_split(uploaded, job / "split", 1))
    else:
        raise ValueError("目前支持 PDF、PNG、JPG、JPEG、WEBP、BMP、TIF 图片文件。")
    def ocr_progress(done, total, _img):
        if progress:
            pct = 50 + int(38 * done / max(total, 1))
            progress(pct, f"OCR 识别中 {done}/{total} 个实际页面")
    pages = ocr_images(split_images, progress_cb=ocr_progress)
    if progress:
        progress(88, "整理 OCR 文本")
    text_path = job / "ocr.txt"
    text = "\n\n".join(f"【第 {page['page']} 页】\n{page['text']}" for page in pages)
    text_path.write_text(text, encoding="utf-8")
    cache_text.write_text(text, encoding="utf-8")
    if progress:
        progress(94, "解析征信字段")
    source_mode = "拍照PDF增强OCR" if suffix == ".pdf" and photo_spread else "OCR识别"
    data = parse_text_result(text_path, uploaded, source_mode)
    if progress:
        progress(100, "生成报告完成")
    return data


INDEX = """<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>征信分析平台</title>
<style>
:root{--ink:#172033;--muted:#66758a;--line:#d8dee8;--panel:#fff;--soft:#f4f7fa;--nav:#172232;--blue:#2864a6;--green:#18794e;--amber:#9a5b00;--red:#b42318;--shadow:0 12px 28px rgba(20,30,45,.08)}
*{box-sizing:border-box} body{margin:0;background:#eef2f5;color:var(--ink);font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","Microsoft YaHei",sans-serif;line-height:1.55} button,input{font:inherit}
.shell{display:grid;grid-template-columns:340px minmax(0,1fr);min-height:100vh}.side{background:var(--nav);color:#fff;padding:22px;display:flex;flex-direction:column;gap:16px}.main{padding:22px;min-width:0}
.brand h1{font-size:22px;margin:0 0 6px}.brand p{margin:0;color:#bfccda;font-size:13px}.upload{border:1px solid rgba(255,255,255,.18);background:rgba(255,255,255,.06);border-radius:8px;padding:14px}.drop{border:1px dashed rgba(255,255,255,.36);border-radius:8px;min-height:118px;display:grid;place-items:center;text-align:center;padding:16px;color:#dce6f1}.drop strong{display:block;color:#fff;margin-bottom:4px}.drop.drag{background:rgba(58,133,214,.18);border-color:#8fc4ff}.native-file{width:100%;margin-top:12px;color:#dce6f1}.file-name{font-size:13px;color:#c5d1df;margin-top:10px;word-break:break-all}.btn-row{display:flex;gap:8px;flex-wrap:wrap;margin-top:12px}button{border:0;border-radius:8px;min-height:40px;padding:10px 13px;cursor:pointer;background:#e8edf3;color:#172033}button.primary{background:#3a85d6;color:#fff}button.ghost{background:transparent;color:#dce6f1;border:1px solid rgba(255,255,255,.24)}button:disabled{opacity:.55;cursor:not-allowed}.note{font-size:13px;color:#bfccda;margin:0}.status{font-size:13px;color:#dce6f1;white-space:pre-wrap;min-height:42px;margin-top:12px}
.steps{display:grid;gap:8px}.step{display:flex;align-items:center;gap:8px;color:#b8c6d6;font-size:13px}.dot{width:9px;height:9px;border-radius:50%;background:#5b6b7f}.step.active{color:#fff}.step.active .dot{background:#5fb3ff}.step.done .dot{background:#4cc38a}
.metrics{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:12px;margin-bottom:16px}.metric{background:var(--panel);border:1px solid var(--line);border-radius:8px;padding:14px;box-shadow:var(--shadow)}.metric span{display:block;color:var(--muted);font-size:13px;margin-bottom:8px}.metric strong{font-size:24px;line-height:1.1;word-break:break-word}.metric small{display:block;color:var(--muted);font-size:12px;margin-top:7px}.panel{background:var(--panel);border:1px solid var(--line);border-radius:8px;box-shadow:var(--shadow);padding:16px;margin-bottom:16px;overflow:hidden}.toolbar{display:flex;justify-content:space-between;align-items:center;gap:12px;margin-bottom:12px}.toolbar h2{font-size:18px;margin:0}.actions{display:flex;gap:8px;flex-wrap:wrap;justify-content:flex-end}.actions button{background:#eef4fb;color:#1f5f9e}.empty{color:var(--muted);padding:26px;text-align:center;background:var(--soft);border:1px dashed var(--line);border-radius:8px}
table{border-collapse:collapse;width:100%;margin:10px 0 18px;min-width:760px}th,td{border:1px solid #cfd7e3;padding:8px;text-align:left;font-size:14px;vertical-align:top}th{background:#f3f6f9;color:#405065}.summary-row th{background:#eaf1f8;color:#1f334a;font-weight:700}.table-wrap{overflow:auto}.report-body h2{font-size:22px;margin:0}.report-body h3{font-size:16px;margin:22px 0 8px}.report-title{display:flex;justify-content:space-between;gap:12px;margin-bottom:12px}.report-body p{margin:7px 0}.report-body ul{margin:8px 0 0 20px;padding:0}.pill{display:inline-flex;align-items:center;border-radius:999px;padding:4px 9px;font-size:13px;font-weight:650;background:#eef4fb;color:var(--blue)}.pill.bad{background:#fdebea;color:var(--red)}.pill.warn{background:#fff4dd;color:var(--amber)}.pill.ok{background:#eaf7f0;color:var(--green)}
.single-report{border-top:1px solid var(--line);margin-top:20px;padding-top:18px;break-before:page}.single-head{display:flex;justify-content:space-between;gap:12px;align-items:flex-start;margin-bottom:8px}.single-head h3{font-size:18px;margin:0}.single-head p{color:var(--muted);font-size:13px;margin:4px 0 0}
.progress{height:10px;background:rgba(255,255,255,.16);border-radius:999px;overflow:hidden;margin-top:10px}.progress span{display:block;height:100%;width:0;background:#5fb3ff;transition:width .25s ease}.debug{font-size:12px;color:var(--muted);background:#f7f9fb;border:1px solid var(--line);border-radius:8px;padding:10px;max-height:120px;overflow:auto;white-space:pre-wrap}
body{min-height:100vh;background:radial-gradient(circle at 18% 10%,rgba(112,231,255,.34),transparent 30%),radial-gradient(circle at 82% 16%,rgba(189,124,255,.30),transparent 32%),linear-gradient(135deg,#0b1328,#15244a 58%,#080d1c);color:#edf7ff}body:before{content:"";position:fixed;inset:0;pointer-events:none;background:linear-gradient(rgba(255,255,255,.035) 1px,transparent 1px),linear-gradient(90deg,rgba(255,255,255,.028) 1px,transparent 1px);background-size:42px 42px;mask-image:radial-gradient(circle at 50% 20%,#000,transparent 75%)}.shell{position:relative}.side{background:linear-gradient(180deg,rgba(255,255,255,.13),rgba(255,255,255,.055));border-right:1px solid rgba(255,255,255,.18);box-shadow:24px 0 70px rgba(0,0,0,.18);backdrop-filter:blur(22px)}.main{background:transparent}.brand h1{letter-spacing:.03em}.brand p,.note,.status,.file-name{color:#bcd0e8}.upload,.metric,.panel,.debug{background:linear-gradient(145deg,rgba(255,255,255,.12),rgba(255,255,255,.055));border:1px solid rgba(255,255,255,.18);border-radius:20px;box-shadow:0 24px 70px rgba(0,0,0,.20);backdrop-filter:blur(20px)}.drop{border-color:rgba(112,231,255,.36);background:rgba(255,255,255,.055);border-radius:18px}.drop.drag{background:rgba(112,231,255,.13);border-color:rgba(112,231,255,.75)}button{border-radius:13px;background:rgba(255,255,255,.13);color:#edf7ff;border:1px solid rgba(255,255,255,.14)}button.primary{background:linear-gradient(135deg,#70e7ff,#7a8cff);color:#06111f;box-shadow:0 12px 30px rgba(112,231,255,.25);font-weight:700}button.ghost{background:rgba(255,255,255,.06);color:#dcecff;border-color:rgba(255,255,255,.20)}button:hover:not(:disabled){transform:translateY(-1px)}.progress{background:rgba(255,255,255,.13)}.progress span{background:linear-gradient(90deg,#70e7ff,#bd7cff)}.metric span,.metric small,.empty,.debug{color:#a9bad2}.metric strong{color:#f1fbff;text-shadow:0 0 22px rgba(112,231,255,.25)}.panel h2,.toolbar h2,.report-body h2,.report-body h3{color:#f3f9ff}.actions button{background:rgba(255,255,255,.10);color:#cff6ff;border:1px solid rgba(112,231,255,.18)}.empty{background:rgba(255,255,255,.055);border-color:rgba(255,255,255,.16);border-radius:18px}.pill{background:rgba(112,231,255,.12);color:#bdf6ff;border:1px solid rgba(112,231,255,.22)}.pill.bad{background:rgba(255,90,118,.14);color:#ffb3c2}.pill.warn{background:rgba(255,214,122,.14);color:#ffe1a3}.pill.ok{background:rgba(78,230,168,.14);color:#b6ffd8}table{color:#eaf4ff}th,td{border-color:rgba(255,255,255,.13)}th{background:rgba(255,255,255,.10);color:#cde1f6}.summary-row th,.summary-row td{background:rgba(112,231,255,.10);color:#f1fbff}.single-report{border-top-color:rgba(255,255,255,.16)}.step{color:#a9bad2}.step.active{color:#f3fbff}.dot{background:rgba(255,255,255,.28)}.step.active .dot{background:#70e7ff;box-shadow:0 0 18px rgba(112,231,255,.8)}.step.done .dot{background:#58efb3}@media(max-width:980px){.shell{grid-template-columns:1fr}.metrics{grid-template-columns:1fr 1fr}.side{min-height:auto}}@media(max-width:620px){.metrics{grid-template-columns:1fr}.main{padding:14px}.side{padding:16px}}@media print{body{background:#fff;color:#172033}.side,.actions,.debug{display:none}.shell{display:block}.main{padding:0}.panel,.metric{box-shadow:none;border:0;background:#fff;color:#172033}.metrics{display:none}table{color:#172033}th,td{border-color:#cfd7e3}th{background:#f3f6f9;color:#405065}}
</style></head><body><div class="shell"><aside class="side"><div class="brand"><h1>征信分析平台</h1><p>扫描件、详版、简版征信均在本机识别和生成报告</p><p><strong>拍照PDF增强版 2026.06.27</strong></p></div>
<section class="upload"><form id="form"><div class="drop" id="drop"><span><strong>拖入多份征信 PDF 或图片</strong><small>同名编号照片（如征信1、征信2）会自动合并为一份报告</small></span></div><input id="file" class="native-file" name="file" type="file" accept=".pdf,application/pdf,image/*" multiple required><div id="fileName" class="file-name">尚未选择文件</div><div class="btn-row"><button id="submitBtn" class="primary">上传并分析</button><button id="resetBtn" class="ghost" type="button">清空结果</button></div></form><div class="progress"><span id="progressBar"></span></div><div id="status" class="status">请选择 PDF 或图片文件，可一次选择多份。</div></section>
<section class="steps"><div class="step" data-step="upload"><span class="dot"></span>上传文件</div><div class="step" data-step="render"><span class="dot"></span>拆页与图像增强</div><div class="step" data-step="ocr"><span class="dot"></span>本地 OCR 识别</div><div class="step" data-step="report"><span class="dot"></span>生成评估报告</div></section><p class="note">扫描件页数多时需要等待。结果仍建议结合原始征信复核关键金额和机构名称。</p></aside>
<main class="main"><section class="metrics"><div class="metric"><span>贷款余额</span><strong id="mLoan">-</strong><small id="mLoanSub">等待分析</small></div><div class="metric"><span>信用卡已用</span><strong id="mCard">-</strong><small id="mCardSub">等待分析</small></div><div class="metric"><span>对外担保余额</span><strong id="mGuarantee">-</strong><small id="mGuaranteeSub">等待分析</small></div><div class="metric"><span>近一年审批</span><strong id="mInquiry">-</strong><small id="mInquirySub">等待分析</small></div></section>
<section class="panel"><div class="toolbar"><h2>分析结果</h2><div class="actions"><button id="downloadPdfBtn" type="button" disabled>下载PDF</button><button id="downloadWordBtn" type="button" disabled>下载Word</button><button id="downloadExcelBtn" type="button" disabled>下载Excel</button><button id="copyBtn" type="button">复制报告</button><button id="printBtn" type="button">打印报告</button></div></div><div id="result"><div class="empty">上传一份或多份征信 PDF/图片后，这里会显示汇总、明细和评估报告。</div></div></section><section class="panel"><div class="toolbar"><h2>识别诊断</h2></div><div id="debug" class="debug">等待分析。</div></section></main></div>
<script>
const form=document.getElementById('form'), fileInput=document.getElementById('file'), drop=document.getElementById('drop'), fileName=document.getElementById('fileName'), statusEl=document.getElementById('status'), result=document.getElementById('result'), debug=document.getElementById('debug'), submitBtn=document.getElementById('submitBtn'), progressBar=document.getElementById('progressBar');
let lastData=null;
const money=v=>Number(v||0).toLocaleString('zh-CN',{maximumFractionDigits:0});
const enterpriseMoney=v=>money(Number(v||0)*10000);
function setSteps(name){const order=['upload','render','ocr','report'];document.querySelectorAll('.step').forEach(el=>{const i=order.indexOf(el.dataset.step),j=order.indexOf(name);el.classList.toggle('done',i<j);el.classList.toggle('active',i===j);});}
function updateFile(){const files=[...fileInput.files];fileName.textContent=files.length?`${files.length} 个文件：`+files.map(f=>f.name).join('，'):'尚未选择文件';}
function setProgress(p){progressBar.style.width=`${Math.max(0,Math.min(100,Number(p)||0))}%`;}
function metricLabel(id,text){document.querySelector(`#${id}`).closest('.metric').querySelector('span').textContent=text}
function setDownloadEnabled(enabled){['downloadPdfBtn','downloadWordBtn','downloadExcelBtn'].forEach(id=>document.getElementById(id).disabled=!enabled);}
function metrics(data){if(data.batch){const a=data.aggregate||{};metricLabel('mLoan','文件数量');document.getElementById('mLoan').textContent=money(a.file_count);document.getElementById('mLoanSub').textContent=`个人 ${a.personal_count||0}，企业 ${a.enterprise_count||0}`;metricLabel('mCard','个人贷款合计');document.getElementById('mCard').textContent=money(a.personal_loan_total);document.getElementById('mCardSub').textContent='单位：元';metricLabel('mGuarantee','企业借贷合计');document.getElementById('mGuarantee').textContent=money(a.enterprise_loan_total);document.getElementById('mGuaranteeSub').textContent='单位：元';metricLabel('mInquiry','风险偏高');document.getElementById('mInquiry').textContent=money(a.high_risk_count);document.getElementById('mInquirySub').textContent='需优先复核';return}const type=(data.report_type||{}).kind;if(type==='enterprise'){const ent=data.enterprise||{}, sum=ent.summary||{}, credit=ent.credit_lines||[], off=ent.off_balance||[], used=credit.reduce((s,x)=>s+Number(x.used||0),0);metricLabel('mLoan','借贷余额');document.getElementById('mLoan').textContent=enterpriseMoney(sum.loan_balance);document.getElementById('mLoanSub').textContent=`${(ent.loans||[]).length} 笔未结清借贷，单位：元`;metricLabel('mCard','授信已用');document.getElementById('mCard').textContent=enterpriseMoney(used);document.getElementById('mCardSub').textContent=`${credit.length} 条授信信息，单位：元`;metricLabel('mGuarantee','表外/担保余额');document.getElementById('mGuarantee').textContent=enterpriseMoney(sum.guarantee_balance);document.getElementById('mGuaranteeSub').textContent=`${off.length} 条表外或担保业务，单位：元`;metricLabel('mInquiry','报告类型');document.getElementById('mInquiry').textContent='企业';document.getElementById('mInquirySub').textContent=(data.report_type||{}).label||'企业征信';return}metricLabel('mLoan','贷款余额');metricLabel('mCard','信用卡已用');metricLabel('mGuarantee','对外担保余额');metricLabel('mInquiry','近一年审批');const loans=(data.loans||[]).filter(x=>Number(x.balance)>0), loanBal=loans.reduce((s,x)=>s+Number(x.balance||0),0), card=data.summary.card_summary||{}, guarantee=data.summary.guarantee_summary||{}, buckets=data.inquiry_buckets||{};document.getElementById('mLoan').textContent=money(loanBal);document.getElementById('mLoanSub').textContent=`${loans.length} 笔有余额贷款`;document.getElementById('mCard').textContent=money(card.used);document.getElementById('mCardSub').textContent=`授信 ${money(card.total_limit)}，使用率 ${card.total_limit?(card.used/card.total_limit*100).toFixed(1):0}%`;document.getElementById('mGuarantee').textContent=money(guarantee.balance);document.getElementById('mGuaranteeSub').textContent=guarantee.balance?'存在对外担保':'未识别到对外担保';const year=(buckets.year||{loan:0,card:0});document.getElementById('mInquiry').textContent=Number(year.loan||0)+Number(year.card||0);document.getElementById('mInquirySub').textContent=`贷款 ${year.loan||0}，信用卡 ${year.card||0}`;}
function render(data){lastData=data;setDownloadEnabled(true);metrics(data);result.innerHTML=data.report_html;if(data.batch){debug.textContent=`识别版本：${data.app_version||'未标注'}\n报告类型：批量征信\n文件数量：${(data.results||[]).length}\n文本总长度：${data.text_length} 字\n个人征信：${(data.aggregate||{}).personal_count||0}\n企业征信：${(data.aggregate||{}).enterprise_count||0}`;return}const type=(data.report_type||{}).label||'未识别';const ent=data.enterprise||{};debug.textContent=`识别版本：${data.app_version||'未标注'}\n报告类型：${type}\n文本长度：${data.text_length} 字\n个人贷款条数：${(data.loans||[]).length}\n信用卡条数：${(data.cards||[]).length}\n企业借贷条数：${(ent.loans||[]).length}\n表外/担保条数：${(ent.off_balance||[]).length}\n文本位置：${data.ocr_text_path||''}\n来源：${data.source_mode||'OCR识别'}`;}
fileInput.addEventListener('change',updateFile);['dragenter','dragover'].forEach(e=>drop.addEventListener(e,ev=>{ev.preventDefault();drop.classList.add('drag')}));['dragleave','drop'].forEach(e=>drop.addEventListener(e,ev=>{ev.preventDefault();drop.classList.remove('drag')}));drop.addEventListener('drop',ev=>{fileInput.files=ev.dataTransfer.files;updateFile();});
async function waitJob(jobId){while(true){const res=await fetch(`/progress?job_id=${encodeURIComponent(jobId)}`);const job=await res.json();if(!res.ok)throw new Error(job.error||'无法读取任务进度');setProgress(job.percent||0);statusEl.textContent=`${job.percent||0}%  ${job.message||'处理中'}`;if((job.message||'').includes('OCR'))setSteps('ocr');else if((job.message||'').includes('渲染')||(job.message||'').includes('图像'))setSteps('render');if(job.status==='done'){setSteps('report');setProgress(100);return job.result}if(job.status==='error'){throw new Error((job.error||job.message||'分析失败')+(job.trace?`\\n${job.trace}`:''))}await new Promise(r=>setTimeout(r,1000));}}
form.addEventListener('submit',async e=>{e.preventDefault();if(!fileInput.files.length){statusEl.textContent='请先选择 PDF 或图片文件。';return}const fd=new FormData(form);submitBtn.disabled=true;lastData=null;setDownloadEnabled(false);setProgress(0);setSteps('upload');statusEl.textContent=`正在上传 ${fileInput.files.length} 个文件...`;result.innerHTML='<div class="empty">正在处理，请稍候。</div>';try{const start=await fetch('/analyze/start',{method:'POST',body:fd});const started=await start.json();if(!start.ok)throw new Error(started.error||'提交任务失败');statusEl.textContent='任务已提交，正在读取进度...';const data=await waitJob(started.job_id);statusEl.textContent=`分析完成：识别 ${data.text_length} 字，生成${data.batch?'批量汇总报告':'报告'}。`;render(data)}catch(err){statusEl.textContent=err.message;debug.textContent=err.stack||String(err)}finally{submitBtn.disabled=false}});
document.getElementById('resetBtn').addEventListener('click',()=>{form.reset();lastData=null;setDownloadEnabled(false);updateFile();setProgress(0);statusEl.textContent='请选择 PDF 或图片文件，可一次选择多份。';result.innerHTML='<div class="empty">上传一份或多份征信 PDF/图片后，这里会显示汇总、明细和评估报告。</div>';debug.textContent='等待分析。';setSteps('')});
function filenameFromDisposition(header, fallback){const match=/filename\\*=UTF-8''([^;]+)/i.exec(header||'');if(match)return decodeURIComponent(match[1]);return fallback;}
async function downloadReport(format){if(!lastData){statusEl.textContent='请先完成分析后再下载。';return}statusEl.textContent='正在生成下载文件...';try{const res=await fetch('/export',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({format,data:lastData})});if(!res.ok){let msg='生成下载文件失败';try{const err=await res.json();msg=err.error||msg}catch(_){msg=await res.text()||msg}throw new Error(msg)}const blob=await res.blob();const ext=format==='pdf'?'pdf':format==='word'?'docx':'xlsx';const name=filenameFromDisposition(res.headers.get('Content-Disposition'),`征信评估报告.${ext}`);const url=URL.createObjectURL(blob);const a=document.createElement('a');a.href=url;a.download=name;document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);statusEl.textContent=`已生成下载文件：${name}`}catch(err){statusEl.textContent=err.message;}}
document.getElementById('downloadPdfBtn').addEventListener('click',()=>downloadReport('pdf'));
document.getElementById('downloadWordBtn').addEventListener('click',()=>downloadReport('word'));
document.getElementById('downloadExcelBtn').addEventListener('click',()=>downloadReport('excel'));
document.getElementById('printBtn').addEventListener('click',()=>window.print());document.getElementById('copyBtn').addEventListener('click',async()=>{await navigator.clipboard.writeText(result.innerText);statusEl.textContent='报告已复制。'});setDownloadEnabled(false);setSteps('');
</script></body></html>"""


class Handler(BaseHTTPRequestHandler):
    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/progress":
            job_id = (parse_qs(parsed.query).get("job_id") or [""])[0]
            job = get_job(job_id)
            if not job:
                self.send_response(404)
                body = json.dumps({"error": "任务不存在"}, ensure_ascii=False).encode("utf-8")
            else:
                self.send_response(200)
                body = json.dumps(job, ensure_ascii=False).encode("utf-8")
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.end_headers()
            self.wfile.write(body)
            return

        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.end_headers()
        self.wfile.write(INDEX.encode("utf-8"))

    def do_POST(self) -> None:
        if self.path == "/export":
            try:
                length = int(self.headers.get("Content-Length") or "0")
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                content, filename, content_type = build_export(payload.get("data") or {}, payload.get("format") or "")
                self.send_response(200)
                self.send_header("Content-Type", content_type)
                self.send_header("Content-Length", str(len(content)))
                self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{quote(filename)}")
                self.end_headers()
                self.wfile.write(content)
            except Exception as exc:
                body = json.dumps({"error": str(exc), "trace": traceback.format_exc()}, ensure_ascii=False).encode("utf-8")
                self.send_response(500)
                self.send_header("Content-Type", "application/json; charset=utf-8")
                self.end_headers()
                self.wfile.write(body)
            return
        if self.path not in {"/analyze", "/analyze/start"}:
            self.send_error(404)
            return
        try:
            form = cgi.FieldStorage(fp=self.rfile, headers=self.headers, environ={"REQUEST_METHOD": "POST"})
            UPLOADS.mkdir(parents=True, exist_ok=True)
            items = form["file"] if isinstance(form["file"], list) else [form["file"]]
            results = []
            for idx, item in enumerate(items, start=1):
                if not getattr(item, "filename", None):
                    continue
                filename = Path(item.filename or f"upload-{idx}.pdf").name
                uploaded = UPLOADS / filename
                if uploaded.exists() and len(items) > 1:
                    uploaded = UPLOADS / f"{uploaded.stem}-{idx}{uploaded.suffix}"
                uploaded.write_bytes(item.file.read())
                results.append(uploaded)
            if not results:
                raise ValueError("未收到可分析的文件。")
            if self.path == "/analyze/start":
                job_id = uuid.uuid4().hex
                set_job(job_id, status="queued", percent=0, message="任务已提交", total_files=len(group_uploads(results)), current_file=0)
                thread = threading.Thread(target=run_analysis_job, args=(job_id, results), daemon=True)
                thread.start()
                data = {"job_id": job_id, "status": "queued"}
            else:
                analyzed = [
                    analyze_image_group(paths) if len(paths) > 1 else analyze_file(paths[0])
                    for paths in group_uploads(results)
                ]
                data = make_batch_response(analyzed) if len(analyzed) > 1 else analyzed[0]
            body = json.dumps(data, ensure_ascii=False).encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.end_headers()
            self.wfile.write(body)
        except Exception as exc:
            body = json.dumps({"error": str(exc), "trace": traceback.format_exc()}, ensure_ascii=False).encode("utf-8")
            self.send_response(500)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.end_headers()
            self.wfile.write(body)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", "8789"))
    server = ThreadingHTTPServer(("127.0.0.1", port), Handler)
    print(f"本地征信 OCR 分析服务（{APP_VERSION}）：http://127.0.0.1:{port}")
    server.serve_forever()
