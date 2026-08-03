from __future__ import annotations

import cgi
import html
import json
import re
import shutil
import sys
import threading
import traceback
import uuid
from collections import defaultdict
from datetime import datetime
from decimal import Decimal
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path, PurePosixPath
from urllib.parse import parse_qs, quote, urlparse


ROOT = Path(__file__).resolve().parents[1]
UPLOADS = ROOT / "tmp" / "comprehensive_uploads"
sys.path.insert(0, str(ROOT / "work"))

import credit_local_server as credit
import flow_local_server as flow


JOBS: dict[str, dict] = {}
JOBS_LOCK = threading.Lock()
FLOW_EXTS = {".xlsx", ".xls", ".csv"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
SUPPORTED_EXTS = FLOW_EXTS | IMAGE_EXTS | {".pdf"}
CREDIT_NAME_RE = re.compile(r"征信|信用报告|信用信息报告|人行报告")
FLOW_NAME_RE = re.compile(r"流水|交易明细|历史明细|账户明细|微信支付|支付宝|对账单|银行明细")
OTHER_NAME_RE = re.compile(r"模板|资产负债|利润表|损益表|财务报表|科目余额|纳税申报|营业执照|合同|申请表")


def set_job(job_id: str, **updates) -> None:
    with JOBS_LOCK:
        JOBS.setdefault(job_id, {}).update(updates)


def get_job(job_id: str) -> dict:
    with JOBS_LOCK:
        return dict(JOBS.get(job_id) or {})


def normalize_text(value) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def money(value) -> str:
    try:
        return f"{float(value or 0):,.2f}"
    except Exception:
        return "0.00"


def safe_filename(value: str) -> str:
    text = re.sub(r'[\\/:*?"<>|\s]+', "_", normalize_text(value) or "客户综合分析报告").strip("_")
    return text[:80] or "客户综合分析报告"


def classify_file(path: Path, original_name: str) -> tuple[str, str]:
    suffix = path.suffix.lower()
    name = original_name or path.name
    if suffix not in SUPPORTED_EXTS:
        return "other", "不支持的文件类型"
    if re.search(r"综合分析报告|征信评估报告|流水统计报告", name):
        return "other", "疑似已生成的报告文件"
    if OTHER_NAME_RE.search(name) and not FLOW_NAME_RE.search(name):
        return "other", "识别为辅助资料，不纳入本次自动统计"
    if suffix in FLOW_EXTS:
        return "flow", "表格文件按流水处理"
    if CREDIT_NAME_RE.search(name):
        return "credit", "文件名识别为征信"
    if FLOW_NAME_RE.search(name):
        return "flow", "文件名识别为流水"
    if suffix == ".pdf":
        return "other", "文件名未包含征信或流水标识，未读取内容"
    if suffix in IMAGE_EXTS:
        return "other", "图片文件名未注明征信或流水，未读取内容"
    return "other", "无法分类"


def prepare_credit_pdf(path: Path, passwords: list[str], target_dir: Path) -> Path:
    if path.suffix.lower() != ".pdf":
        return path
    try:
        from pypdf import PdfReader, PdfWriter

        reader = PdfReader(str(path))
        if not reader.is_encrypted:
            return path
        unlocked = False
        for password in passwords or [""]:
            try:
                if reader.decrypt(password or ""):
                    unlocked = True
                    break
            except Exception:
                continue
        if not unlocked:
            raise ValueError(f"{path.name} 需要正确的 PDF 密码。")
        target = target_dir / f"{path.stem}-已解密.pdf"
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        with target.open("wb") as handle:
            writer.write(handle)
        return target
    except ValueError:
        raise
    except Exception:
        return path


def source_filename(value: str) -> str:
    return normalize_text(value).split(" / ", 1)[0]


def rewrite_flow_sources(result: dict, saved_name: str, original_name: str) -> None:
    result["filename"] = original_name
    for txn in result.get("transactions") or []:
        source = normalize_text(txn.get("source"))
        if source == saved_name:
            txn["source"] = original_name
        elif source.startswith(saved_name + " / "):
            txn["source"] = original_name + source[len(saved_name) :]


def flow_monthly(transactions: list[dict]) -> list[dict]:
    months: dict[str, dict] = defaultdict(lambda: {"income": 0.0, "expense": 0.0, "count": 0})
    for txn in transactions:
        if txn.get("default_excluded"):
            continue
        month = txn.get("month") or "未识别日期"
        months[month]["income"] += float(txn.get("income") or 0)
        months[month]["expense"] += float(txn.get("expense") or 0)
        months[month]["count"] += 1
    return [
        {"month": key, "income": round(value["income"], 2), "expense": round(value["expense"], 2), "count": value["count"]}
        for key, value in sorted(months.items())
    ]


def flow_accounts(transactions: list[dict]) -> list[dict]:
    accounts: dict[str, dict] = defaultdict(lambda: {"income": 0.0, "expense": 0.0, "count": 0, "months": set(), "files": set()})
    for txn in transactions:
        if txn.get("default_excluded"):
            continue
        account = txn.get("account") or "未识别账户"
        item = accounts[account]
        item["income"] += float(txn.get("income") or 0)
        item["expense"] += float(txn.get("expense") or 0)
        item["count"] += 1
        item["months"].add(txn.get("month") or "未识别日期")
        item["files"].add(source_filename(txn.get("source") or ""))
    result = []
    for account, item in accounts.items():
        result.append({
            "account": account,
            "income": round(item["income"], 2),
            "expense": round(item["expense"], 2),
            "count": item["count"],
            "month_count": len(item["months"]),
            "file_count": len(item["files"]),
        })
    return sorted(result, key=lambda item: item["income"], reverse=True)


def top_counterparties(transactions: list[dict], limit: int = 20) -> list[dict]:
    grouped: dict[str, dict] = defaultdict(lambda: {"income": 0.0, "expense": 0.0, "count": 0})
    for txn in transactions:
        if txn.get("default_excluded"):
            continue
        name = txn.get("counterparty") or "未识别"
        grouped[name]["income"] += float(txn.get("income") or 0)
        grouped[name]["expense"] += float(txn.get("expense") or 0)
        grouped[name]["count"] += 1
    rows = [{"name": key, **value} for key, value in grouped.items()]
    return sorted(rows, key=lambda item: item["income"], reverse=True)[:limit]


def table_block(title: str, headers: list[str], rows: list[list]) -> dict:
    return {"title": title, "headers": headers, "rows": rows or [["无数据"] + [""] * (len(headers) - 1)]}


def build_comprehensive_result(
    folder_name: str,
    file_records: list[dict],
    credit_results: list[dict],
    flow_response: dict,
    errors: list[dict],
) -> dict:
    credit_summaries = [credit.result_summary(item) for item in credit_results]
    transactions = list(flow_response.get("transactions") or [])
    monthly = flow_monthly(transactions)
    accounts = flow_accounts(transactions)
    counterparties = top_counterparties(transactions)
    valid_income = round(sum(item["income"] for item in monthly), 2)
    valid_expense = round(sum(item["expense"] for item in monthly), 2)
    month_count = sum(1 for item in monthly if item["count"] > 0)
    avg_income = round(valid_income / month_count, 2) if month_count else 0

    personal_loan = sum(float(item.get("personal_loan_balance") or 0) for item in credit_summaries)
    enterprise_loan = sum(float(item.get("enterprise_loan_balance") or 0) for item in credit_summaries)
    loan_balance = round(personal_loan + enterprise_loan, 2)
    card_used = round(sum(float(item.get("card_used") or 0) for item in credit_summaries), 2)
    guarantee_balance = round(sum(float(item.get("personal_guarantee_balance") or 0) + float(item.get("enterprise_guarantee_balance") or 0) for item in credit_summaries), 2)
    overdue_total = round(sum(float(item.get("overdue_total") or 0) for item in credit_summaries), 2)
    high_risk_count = sum(item.get("risk") == "风险偏高" for item in credit_summaries)
    coverage_ratio = round(valid_income / loan_balance, 2) if loan_balance else None
    credit_count = len(credit_results)
    flow_file_count = len(flow_response.get("files") or [])
    incomplete = credit_count == 0 or flow_file_count == 0
    if high_risk_count or overdue_total > 0:
        risk = "需重点复核"
    elif incomplete:
        risk = "资料不完整"
    else:
        risk = "一般关注"

    warnings = []
    if high_risk_count:
        warnings.append(f"{high_risk_count} 份征信被识别为风险偏高，应优先复核逾期、担保和高使用率授信。")
    if overdue_total:
        warnings.append(f"企业征信识别逾期总额 {money(overdue_total)} 元。")
    if guarantee_balance:
        warnings.append(f"识别对外担保或表外余额 {money(guarantee_balance)} 元，建议核实实际代偿风险。")
    if flow_response.get("duplicate_transaction_count"):
        warnings.append(f"跨文件自动去除 {flow_response['duplicate_transaction_count']} 笔重复流水。")
    if incomplete:
        warnings.append("征信或流水资料缺失，当前综合判断仅供阶段性参考。")
    if not warnings:
        warnings.append("未触发高风险规则，仍应结合原始征信、流水用途及客户经营情况人工复核。")

    classification_rows = [
        [item.get("relative_path") or item.get("name"), item.get("category_label"), item.get("reason"), item.get("status")]
        for item in file_records
    ]
    credit_template_rows = []
    credit_summary_iter = iter(credit_summaries)
    for item in [record for record in file_records if record.get("category") == "credit"]:
        summary = {}
        if item.get("status") == "分析完成":
            summary = next(credit_summary_iter, {}) or {}
        credit_template_rows.append([
            item.get("relative_path") or item.get("name"),
            summary.get("name") or "",
            summary.get("type") or "",
            summary.get("report_date") or "未识别",
            summary.get("risk") or "",
            item.get("status") or "",
        ])
    flow_file_map = {normalize_text(item.get("filename")): item for item in flow_response.get("files") or []}
    flow_template_rows = []
    for item in [record for record in file_records if record.get("category") == "flow"]:
        flow_file = flow_file_map.get(normalize_text(item.get("name"))) or {}
        accounts_text = "、".join(flow_file.get("accounts") or []) or "未识别"
        flow_template_rows.append([
            item.get("relative_path") or item.get("name"),
            accounts_text,
            flow_file.get("source_mode") or "",
            flow_file.get("transaction_count") or 0,
            flow_file.get("unique_transaction_count") if flow_file.get("unique_transaction_count") is not None else flow_file.get("transaction_count") or 0,
            item.get("status") or "",
        ])
    other_template_rows = [
        [item.get("relative_path") or item.get("name"), item.get("reason"), item.get("status")]
        for item in file_records
        if item.get("category") == "other"
    ]
    credit_rows = [[
        item.get("name"), item.get("type"), item.get("report_date") or "未识别", item.get("risk"),
        item.get("loan_count") or 0, item.get("personal_loan_balance") or 0, item.get("card_used") or 0,
        item.get("enterprise_loan_balance") or 0, item.get("personal_guarantee_balance") or 0,
        item.get("enterprise_guarantee_balance") or 0, item.get("overdue_total") or 0,
    ] for item in credit_summaries]
    account_rows = [[item["account"], item["file_count"], item["month_count"], item["count"], item["income"], item["expense"]] for item in accounts]
    monthly_rows = [[item["month"], item["count"], item["income"], item["expense"], round(item["income"] - item["expense"], 2)] for item in monthly]
    counterparty_rows = [[item["name"], item["count"], round(item["income"], 2), round(item["expense"], 2)] for item in counterparties]

    blocks = [
        table_block("综合结论", ["项目", "数值"], [
            ["客户文件夹", folder_name or "未命名文件夹"],
            ["综合判断", risk],
            ["征信文件", credit_count],
            ["流水文件", flow_file_count],
            ["流水唯一交易", len(transactions)],
            ["跨文件重复交易", flow_response.get("duplicate_transaction_count") or 0],
            ["贷款余额", loan_balance],
            ["信用卡已用", card_used],
            ["担保/表外余额", guarantee_balance],
            ["有效收入流水", valid_income],
            ["月均有效收入", avg_income],
            ["收入流水/贷款余额（参考）", coverage_ratio if coverage_ratio is not None else "不适用"],
        ]),
        table_block("征信资料模板", ["文件", "客户/企业", "类型", "报告日期", "风险", "状态"], credit_template_rows),
        table_block("流水资料模板", ["文件", "识别账户", "读取方式", "识别笔数", "去重后笔数", "状态"], flow_template_rows),
        table_block("其他/跳过文件", ["文件", "判断依据", "状态"], other_template_rows),
        table_block("文件自动分类", ["文件", "分类", "判断依据", "状态"], classification_rows),
        table_block("征信汇总", ["客户/企业", "类型", "报告日期", "风险", "未结清笔数", "个人贷款余额", "信用卡已用", "企业借贷余额", "个人担保", "企业表外/担保", "逾期总额"], credit_rows),
        table_block("流水账户汇总", ["账户", "文件数", "月份数", "交易笔数", "有效收入", "有效支出"], account_rows),
        table_block("流水月度汇总", ["月份", "有效笔数", "有效收入", "有效支出", "净流水"], monthly_rows),
        table_block("主要收入交易对象", ["交易对象", "笔数", "收入", "支出"], counterparty_rows),
        table_block("风险提示与复核建议", ["序号", "提示"], [[idx, text] for idx, text in enumerate(warnings, start=1)]),
        table_block("处理异常", ["文件", "问题"], [[item.get("file"), item.get("error")] for item in errors]),
    ]
    for idx, item in enumerate(credit_results, start=1):
        name = credit.result_summary(item).get("name") or f"第{idx}份征信"
        for block in credit.result_export_blocks(item):
            blocks.append({**block, "title": f"征信明细 - {name} - {block['title']}"})

    def td(value):
        return html.escape(str(value if value is not None else ""))

    credit_html = "".join(
        f"<tr><td>{td(row[0])}</td><td>{td(row[1])}</td><td>{td(row[2])}</td><td>{td(row[3])}</td><td>{row[4]}</td><td>{money(row[5])}</td><td>{money(row[6])}</td><td>{money(row[7])}</td><td>{money(row[8] + row[9])}</td><td>{money(row[10])}</td></tr>"
        for row in credit_rows
    ) or '<tr><td colspan="10">未识别到征信文件</td></tr>'
    account_html = "".join(
        f"<tr><td>{td(row[0])}</td><td>{row[1]}</td><td>{row[2]}</td><td>{row[3]}</td><td>{money(row[4])}</td><td>{money(row[5])}</td></tr>"
        for row in account_rows
    ) or '<tr><td colspan="6">未识别到流水账户</td></tr>'
    month_html = "".join(
        f"<tr><td>{td(row[0])}</td><td>{row[1]}</td><td>{money(row[2])}</td><td>{money(row[3])}</td><td>{money(row[4])}</td></tr>"
        for row in monthly_rows
    ) or '<tr><td colspan="5">暂无月度流水</td></tr>'
    file_html = "".join(
        f"<tr><td>{td(row[0])}</td><td>{td(row[1])}</td><td>{td(row[2])}</td><td>{td(row[3])}</td></tr>"
        for row in classification_rows
    )
    credit_template_html = "".join(
        f"<tr><td>{td(row[0])}</td><td>{td(row[1])}</td><td>{td(row[2])}</td><td>{td(row[3])}</td><td>{td(row[4])}</td><td>{td(row[5])}</td></tr>"
        for row in credit_template_rows
    ) or '<tr><td colspan="6">未识别到征信资料</td></tr>'
    flow_template_html = "".join(
        f"<tr><td>{td(row[0])}</td><td>{td(row[1])}</td><td>{td(row[2])}</td><td>{row[3]}</td><td>{row[4]}</td><td>{td(row[5])}</td></tr>"
        for row in flow_template_rows
    ) or '<tr><td colspan="6">未识别到流水资料</td></tr>'
    other_template_html = "".join(
        f"<tr><td>{td(row[0])}</td><td>{td(row[1])}</td><td>{td(row[2])}</td></tr>"
        for row in other_template_rows
    ) or '<tr><td colspan="3">暂无其他/跳过文件</td></tr>'
    warning_html = "".join(f"<li>{td(item)}</li>" for item in warnings)
    detail_html = "".join(
        f'<section class="credit-detail"><h3>征信明细：{td(credit.result_summary(item).get("name") or f"第{idx}份")}</h3>{item.get("report_html") or ""}</section>'
        for idx, item in enumerate(credit_results, start=1)
    )
    report_html = f"""
    <article class="report-body">
      <div class="report-title"><div><h2>{td(folder_name or '客户')}综合分析报告</h2><p>生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}；综合判断：<strong>{td(risk)}</strong></p></div><span class="pill">{td(risk)}</span></div>
      <h3>一、综合指标</h3>
      <table><thead><tr><th>贷款余额</th><th>信用卡已用</th><th>担保/表外余额</th><th>有效收入流水</th><th>月均有效收入</th><th>流水月份</th></tr></thead><tbody><tr><td>{money(loan_balance)}</td><td>{money(card_used)}</td><td>{money(guarantee_balance)}</td><td>{money(valid_income)}</td><td>{money(avg_income)}</td><td>{month_count}</td></tr></tbody></table>
      <h3>二、风险提示与复核建议</h3><ul>{warning_html}</ul>
      <h3>三、资料分类模板</h3>
      <div class="class-grid">
        <section class="class-card"><h4>征信资料模板</h4><p class="note">用于查看个人/企业征信文件是否识别成功，以及报告主体、日期和风险等级。</p><div class="table-wrap"><table><thead><tr><th>文件</th><th>客户/企业</th><th>类型</th><th>报告日期</th><th>风险</th><th>状态</th></tr></thead><tbody>{credit_template_html}</tbody></table></div></section>
        <section class="class-card"><h4>流水资料模板</h4><p class="note">用于查看银行流水、微信/支付宝明细、Excel 流水等是否进入流水统计。</p><div class="table-wrap"><table><thead><tr><th>文件</th><th>识别账户</th><th>读取方式</th><th>识别笔数</th><th>去重后笔数</th><th>状态</th></tr></thead><tbody>{flow_template_html}</tbody></table></div></section>
      </div>
      <details class="other-files"><summary>其他/跳过文件</summary><table><thead><tr><th>文件</th><th>判断依据</th><th>状态</th></tr></thead><tbody>{other_template_html}</tbody></table></details>
      <h3>四、文件自动分类明细</h3><table><thead><tr><th>文件</th><th>分类</th><th>判断依据</th><th>状态</th></tr></thead><tbody>{file_html}</tbody></table>
      <h3>五、征信汇总</h3><table><thead><tr><th>客户/企业</th><th>类型</th><th>报告日期</th><th>风险</th><th>未结清笔数</th><th>个人贷款</th><th>信用卡已用</th><th>企业借贷</th><th>担保/表外</th><th>逾期总额</th></tr></thead><tbody>{credit_html}</tbody></table>
      <h3>六、流水账户汇总</h3><table><thead><tr><th>账户</th><th>文件数</th><th>月份数</th><th>有效笔数</th><th>有效收入</th><th>有效支出</th></tr></thead><tbody>{account_html}</tbody></table>
      <h3>七、流水月度汇总</h3><table><thead><tr><th>月份</th><th>有效笔数</th><th>有效收入</th><th>有效支出</th><th>净流水</th></tr></thead><tbody>{month_html}</tbody></table>
      <p class="note">本报告由本机自动识别生成。自动分类、OCR 字段、关键金额、交易性质和综合判断均建议结合原件及人工尽调复核。</p>
      {detail_html}
    </article>
    """
    return {
        "folder_name": folder_name,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "risk": risk,
        "report_html": report_html,
        "export_blocks": blocks,
        "metrics": {
            "loan_balance": loan_balance,
            "card_used": card_used,
            "guarantee_balance": guarantee_balance,
            "valid_income": valid_income,
            "valid_expense": valid_expense,
            "avg_income": avg_income,
            "month_count": month_count,
            "credit_count": credit_count,
            "flow_file_count": flow_file_count,
            "transaction_count": len(transactions),
            "duplicate_count": flow_response.get("duplicate_transaction_count") or 0,
        },
        "files": file_records,
        "errors": errors,
    }


def run_comprehensive_job(job_id: str, manifest: list[dict], folder_name: str, passwords: list[str]) -> None:
    try:
        credit_results = []
        flow_results = []
        flow_transactions = []
        errors = []
        total = len(manifest) or 1
        for idx, item in enumerate(manifest, start=1):
            path = Path(item["path"])
            category, reason = classify_file(path, item["name"])
            item["category"] = category
            item["category_label"] = {"credit": "征信", "flow": "流水", "other": "其他/跳过"}[category]
            item["reason"] = reason
            if category == "other":
                item["status"] = "已跳过"
                continue

            def progress(file_pct, message, idx=idx, item=item):
                overall = min(99, int(((idx - 1) + file_pct / 100) * 100 / total))
                set_job(job_id, status="running", percent=overall, message=f"第 {idx}/{total} 个文件：{item['name']} - {message}")

            try:
                if category == "credit":
                    analysis_path = prepare_credit_pdf(path, passwords, path.parent)
                    result = credit.analyze_file(analysis_path, progress=progress)
                    result["filename"] = item["name"]
                    credit_results.append(result)
                else:
                    result = flow.analyze_file(path, progress=progress, passwords=passwords)
                    rewrite_flow_sources(result, path.name, item["name"])
                    flow_results.append({key: value for key, value in result.items() if key != "transactions"})
                    flow_transactions.extend(result.get("transactions") or [])
                item["status"] = "分析完成"
            except Exception as exc:
                item["status"] = "分析失败"
                errors.append({"file": item["relative_path"], "error": str(exc)})

        flow_response = flow.build_flow_response(flow_results, flow_transactions)
        result = build_comprehensive_result(folder_name, manifest, credit_results, flow_response, errors)
        set_job(job_id, status="done", percent=100, message="综合分析完成", result=result)
    except Exception as exc:
        set_job(job_id, status="error", percent=100, message=str(exc), error=str(exc), trace=traceback.format_exc())


def build_docx(result: dict) -> bytes:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(9)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{result.get('folder_name') or '客户'}综合分析报告")
    run.bold = True
    run.font.size = Pt(18)
    for block in result.get("export_blocks") or []:
        heading = doc.add_paragraph()
        heading.add_run(block["title"]).bold = True
        table = doc.add_table(rows=1, cols=len(block["headers"]))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, header in zip(table.rows[0].cells, block["headers"]):
            cell.text = str(header)
            cell.paragraphs[0].runs[0].bold = True
        for row in block["rows"]:
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                cell.text = money(value) if isinstance(value, (int, float, Decimal)) and not isinstance(value, bool) else str(value or "")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_xlsx(result: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)
    used = set()
    for block in result.get("export_blocks") or []:
        base = re.sub(r"[\[\]:*?/\\]", "", block["title"])[:28] or "报告"
        name = base
        idx = 1
        while name in used:
            name = f"{base[:27]}{idx}"
            idx += 1
        used.add(name)
        ws = wb.create_sheet(name)
        ws.append(block["headers"])
        for row in block["rows"]:
            ws.append(row)
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.font = Font(bold=True, color="1F334A")
            cell.fill = PatternFill("solid", fgColor="EAF1F8")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                if isinstance(cell.value, (int, float)):
                    cell.number_format = "#,##0.00"
        for col in range(1, ws.max_column + 1):
            letter = get_column_letter(col)
            width = max(len(str(ws.cell(row=row, column=col).value or "")) for row in range(1, min(ws.max_row, 100) + 1)) + 3
            ws.column_dimensions[letter].width = min(max(width, 12), 36)
    bio = BytesIO()
    wb.save(bio)
    return bio.getvalue()


def build_pdf(result: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = "Helvetica"
    for font_path in [
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]:
        try:
            pdfmetrics.registerFont(TTFont("ComprehensiveCN", font_path))
            font_name = "ComprehensiveCN"
            break
        except Exception:
            continue
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCN", parent=styles["Title"], fontName=font_name, fontSize=18, leading=24, alignment=1)
    heading_style = ParagraphStyle("HeadingCN", parent=styles["Heading2"], fontName=font_name, fontSize=12, leading=16, spaceBefore=8, spaceAfter=4)
    cell_style = ParagraphStyle("CellCN", parent=styles["BodyText"], fontName=font_name, fontSize=6.5, leading=8)
    bio = BytesIO()
    page = landscape(A4)
    doc = SimpleDocTemplate(bio, pagesize=page, leftMargin=9 * mm, rightMargin=9 * mm, topMargin=9 * mm, bottomMargin=9 * mm)
    story = [Paragraph(html.escape(f"{result.get('folder_name') or '客户'}综合分析报告"), title_style), Spacer(1, 4 * mm)]
    usable = page[0] - 18 * mm
    for block in result.get("export_blocks") or []:
        story.append(Paragraph(html.escape(block["title"]), heading_style))
        rows = [[Paragraph(html.escape(str(value)), cell_style) for value in block["headers"]]]
        for row in block["rows"]:
            rows.append([Paragraph(html.escape(money(value) if isinstance(value, (int, float, Decimal)) and not isinstance(value, bool) else str(value or "")), cell_style) for value in row])
        widths = [usable / max(1, len(block["headers"]))] * len(block["headers"])
        table = Table(rows, colWidths=widths, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF1F8")),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 2.5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 2.5),
            ("TOPPADDING", (0, 0), (-1, -1), 2.5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
        ]))
        story.extend([table, Spacer(1, 3 * mm)])
    doc.build(story)
    return bio.getvalue()


def build_export(result: dict, fmt: str) -> tuple[bytes, str, str]:
    title = safe_filename(f"{result.get('folder_name') or '客户'}综合分析报告")
    fmt = normalize_text(fmt).lower()
    if fmt == "pdf":
        return build_pdf(result), f"{title}.pdf", "application/pdf"
    if fmt in {"word", "docx"}:
        return build_docx(result), f"{title}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt in {"excel", "xlsx"}:
        return build_xlsx(result), f"{title}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    raise ValueError("不支持的下载格式。")


INDEX = r"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>征信与流水综合分析</title><style>
:root{--ink:#edf7ff;--muted:#a9bad2;--line:rgba(255,255,255,.16);--panel:rgba(255,255,255,.10);--nav:transparent;--blue:#70e7ff;--green:#58efb3;--red:#ff6b8a;--shadow:0 24px 70px rgba(0,0,0,.20)}*{box-sizing:border-box}body{margin:0;min-height:100vh;background:radial-gradient(circle at 18% 10%,rgba(112,231,255,.34),transparent 30%),radial-gradient(circle at 82% 16%,rgba(189,124,255,.30),transparent 32%),linear-gradient(135deg,#0b1328,#15244a 58%,#080d1c);color:var(--ink);font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","Microsoft YaHei",sans-serif;line-height:1.5}body:before{content:"";position:fixed;inset:0;pointer-events:none;background:linear-gradient(rgba(255,255,255,.035) 1px,transparent 1px),linear-gradient(90deg,rgba(255,255,255,.028) 1px,transparent 1px);background-size:42px 42px;mask-image:radial-gradient(circle at 50% 20%,#000,transparent 75%)}.shell{position:relative;display:grid;grid-template-columns:360px minmax(0,1fr);min-height:100vh}.side{background:linear-gradient(180deg,rgba(255,255,255,.13),rgba(255,255,255,.055));border-right:1px solid rgba(255,255,255,.18);box-shadow:24px 0 70px rgba(0,0,0,.18);backdrop-filter:blur(22px);color:#fff;padding:22px}.main{padding:22px;min-width:0;background:transparent}.brand h1{font-size:22px;margin:0 0 6px;letter-spacing:.03em}.brand p{font-size:13px;color:#bcd0e8;margin:0 0 16px}.box{border:1px solid rgba(255,255,255,.18);background:linear-gradient(145deg,rgba(255,255,255,.12),rgba(255,255,255,.055));border-radius:20px;padding:14px;margin-bottom:14px;box-shadow:0 24px 70px rgba(0,0,0,.20);backdrop-filter:blur(20px)}.template-mini h3{font-size:15px;margin:0 0 10px}.template-mini ul{margin:8px 0 0 18px;padding:0;color:#dce6f1;font-size:13px}.template-mini li{margin:5px 0}.drop{border:1px dashed rgba(112,231,255,.36);border-radius:18px;padding:18px;text-align:center;color:#dce6f1;background:rgba(255,255,255,.055)}.drop strong{display:block;color:#fff}.native-file,textarea,input{width:100%;margin-top:10px}textarea,input{border:1px solid rgba(255,255,255,.18);background:rgba(8,15,34,.46);color:#edf7ff;border-radius:14px;padding:9px}textarea{min-height:72px}.file-name,.status{font-size:13px;color:#dce6f1;margin-top:10px;white-space:pre-wrap;word-break:break-all}.btn-row{display:flex;gap:8px;margin-top:12px;flex-wrap:wrap}button{border:1px solid rgba(255,255,255,.14);border-radius:13px;padding:10px 14px;cursor:pointer;background:rgba(255,255,255,.13);color:#edf7ff}button.primary{background:linear-gradient(135deg,#70e7ff,#7a8cff);color:#06111f;box-shadow:0 12px 30px rgba(112,231,255,.25);font-weight:700}button.ghost{background:rgba(255,255,255,.06);border-color:rgba(255,255,255,.20);color:#dcecff}button:disabled{opacity:.55}.progress{height:10px;background:rgba(255,255,255,.13);border-radius:999px;overflow:hidden;margin-top:10px}.progress span{display:block;height:100%;width:0;background:linear-gradient(90deg,#70e7ff,#bd7cff)}.metrics{display:grid;grid-template-columns:repeat(4,minmax(0,1fr));gap:12px;margin-bottom:16px}.metric,.panel,.class-card,.other-files{background:linear-gradient(145deg,rgba(255,255,255,.12),rgba(255,255,255,.055));border:1px solid rgba(255,255,255,.18);border-radius:20px;box-shadow:0 24px 70px rgba(0,0,0,.20);backdrop-filter:blur(20px)}.metric{padding:14px}.metric span{display:block;color:#a9bad2;font-size:13px}.metric strong{display:block;font-size:24px;margin-top:8px;color:#f1fbff;text-shadow:0 0 22px rgba(112,231,255,.25)}.metric small{color:#a9bad2}.panel{padding:16px;margin-bottom:16px;overflow:hidden}.toolbar{display:flex;justify-content:space-between;gap:12px;align-items:center;margin-bottom:12px}.toolbar h2{font-size:18px;margin:0;color:#f3f9ff}.actions{display:flex;gap:8px;flex-wrap:wrap}.actions button{background:rgba(255,255,255,.10);color:#cff6ff;border:1px solid rgba(112,231,255,.18)}.empty{color:#a9bad2;background:rgba(255,255,255,.055);border:1px dashed rgba(255,255,255,.16);padding:28px;text-align:center;border-radius:18px}table{border-collapse:collapse;width:100%;min-width:820px;margin:10px 0 18px;color:#eaf4ff}th,td{border:1px solid rgba(255,255,255,.13);padding:8px;text-align:left;font-size:14px;vertical-align:top}th{background:rgba(255,255,255,.10);color:#cde1f6}.table-wrap{overflow:auto}.report-body h2{font-size:22px;margin:0;color:#f3f9ff}.report-body h3{font-size:16px;margin:22px 0 8px;color:#f3f9ff}.report-body h4{font-size:15px;margin:0 0 4px;color:#f3f9ff}.report-title{display:flex;justify-content:space-between;gap:12px}.pill{display:inline-flex;height:fit-content;border-radius:999px;padding:4px 10px;background:rgba(112,231,255,.12);color:#bdf6ff;border:1px solid rgba(112,231,255,.22)}.note{font-size:13px;color:#a9bad2}.class-grid{display:grid;grid-template-columns:1fr 1fr;gap:12px;margin:10px 0 12px}.class-card{padding:12px;overflow:hidden}.class-card table{min-width:680px;margin-bottom:4px}.other-files{border-style:dashed;padding:10px 12px;margin-bottom:14px}.other-files summary{cursor:pointer;color:#d9efff;font-weight:600}.credit-detail{border-top:2px solid rgba(255,255,255,.16);margin-top:28px;padding-top:16px}.credit-detail .report-title{display:none}.debug{white-space:pre-wrap;font-size:12px;color:#a9bad2;max-height:140px;overflow:auto}@media(max-width:980px){.shell{grid-template-columns:1fr}.metrics{grid-template-columns:1fr 1fr}.class-grid{grid-template-columns:1fr}}@media(max-width:620px){.metrics{grid-template-columns:1fr}.main,.side{padding:14px}}@media print{body{background:#fff;color:#172033}.side,.actions,.debug{display:none}.shell{display:block}.main{padding:0}.panel,.metric,.class-card{box-shadow:none;background:#fff;color:#172033;border:0}table{color:#172033}th,td{border-color:#cfd7e3}th{background:#f3f6f9;color:#405065}}
</style></head><body><div class="shell"><aside class="side"><div class="brand"><h1>客户综合分析</h1><p>选择一个同时包含征信和流水的客户文件夹，系统将在本机自动分类、分析并生成综合报告。</p></div>
<section class="box"><form id="form"><div class="drop" id="drop"><strong>选择客户文件夹</strong><small>支持征信 PDF/图片与流水 PDF/Excel/图片</small></div><input id="folder" class="native-file" type="file" webkitdirectory directory multiple required><div id="fileName" class="file-name">尚未选择文件夹</div><label><small>文件密码（如有，每行一个）</small><textarea id="passwords" placeholder="可填写 PDF 或流水文件密码"></textarea></label><div class="btn-row"><button id="submitBtn" class="primary">自动分类并综合分析</button><button id="resetBtn" class="ghost" type="button">清空</button></div></form><div class="progress"><span id="progressBar"></span></div><div id="status" class="status">请选择客户文件夹。</div></section><section class="box template-mini"><h3>分类查看模板</h3><ul><li><strong>征信资料：</strong>文件名含“征信、信用报告、人行报告”等，会进入征信汇总。</li><li><strong>流水资料：</strong>PDF/Excel/图片流水、交易明细、历史明细、微信/支付宝明细，会进入流水统计。</li><li><strong>其他资料：</strong>营业执照、合同、财报等默认跳过，不影响本次综合分析。</li></ul></section><p class="note" style="color:#c7d1df">无法可靠分类的文件会列为“其他/跳过”，不会影响其余文件。综合结论仅作辅助，关键数据仍需结合原件复核。</p></aside>
<main class="main"><section class="metrics"><div class="metric"><span>综合判断</span><strong id="mRisk">-</strong><small>等待分析</small></div><div class="metric"><span>贷款余额</span><strong id="mLoan">-</strong><small>征信口径</small></div><div class="metric"><span>有效收入流水</span><strong id="mIncome">-</strong><small>自动剔除常见非经营项</small></div><div class="metric"><span>月均有效收入</span><strong id="mAvg">-</strong><small>按有效月份</small></div></section>
<section class="panel"><div class="toolbar"><h2>综合报告</h2><div class="actions"><button id="pdfBtn" disabled>下载PDF</button><button id="wordBtn" disabled>下载Word</button><button id="excelBtn" disabled>下载Excel</button><button id="printBtn">打印</button></div></div><div id="result"><div class="empty">选择客户文件夹后，这里会显示征信、流水与综合风险提示。</div></div></section><section class="panel"><h2>处理诊断</h2><div id="debug" class="debug">等待分析。</div></section></main></div>
<script>
const folder=document.getElementById('folder'),fileName=document.getElementById('fileName'),statusEl=document.getElementById('status'),progressBar=document.getElementById('progressBar'),result=document.getElementById('result'),debug=document.getElementById('debug'),submitBtn=document.getElementById('submitBtn');let currentJobId=null,lastData=null;
const money=v=>Number(v||0).toLocaleString('zh-CN',{minimumFractionDigits:2,maximumFractionDigits:2});
function rootName(){const f=folder.files[0];return f&&f.webkitRelativePath?f.webkitRelativePath.split('/')[0]:'客户文件夹'}
function updateFiles(){const files=[...folder.files];fileName.textContent=files.length?`${rootName()}：${files.length} 个文件`:'尚未选择文件夹'}
function setProgress(v){progressBar.style.width=`${Math.max(0,Math.min(100,Number(v)||0))}%`}
function setDownloads(on){['pdfBtn','wordBtn','excelBtn'].forEach(id=>document.getElementById(id).disabled=!on)}
folder.addEventListener('change',updateFiles);
document.getElementById('form').addEventListener('submit',async e=>{e.preventDefault();const files=[...folder.files];if(!files.length){statusEl.textContent='请先选择客户文件夹。';return}const fd=new FormData();for(const file of files)fd.append('file',file,file.webkitRelativePath||file.name);fd.append('folder_name',rootName());fd.append('passwords',document.getElementById('passwords').value||'');submitBtn.disabled=true;setDownloads(false);setProgress(0);result.innerHTML='<div class="empty">正在自动分类并分析，请稍候。</div>';try{const start=await fetch('/analyze/start',{method:'POST',body:fd});const started=await start.json();if(!start.ok)throw new Error(started.error||'提交失败');currentJobId=started.job_id;while(true){const res=await fetch(`/progress?job_id=${encodeURIComponent(currentJobId)}`);const job=await res.json();if(!res.ok)throw new Error(job.error||'读取进度失败');setProgress(job.percent||0);statusEl.textContent=`${job.percent||0}%  ${job.message||'处理中'}`;if(job.status==='done'){lastData=job.result;setDownloads(true);setProgress(100);render(job.result);statusEl.textContent='综合分析完成。';break}if(job.status==='error')throw new Error(job.error||job.message||'分析失败');await new Promise(r=>setTimeout(r,1000))}}catch(err){statusEl.textContent=err.message;debug.textContent=err.stack||String(err)}finally{submitBtn.disabled=false}});
function render(data){const m=data.metrics||{};document.getElementById('mRisk').textContent=data.risk||'-';document.getElementById('mLoan').textContent=money(m.loan_balance);document.getElementById('mIncome').textContent=money(m.valid_income);document.getElementById('mAvg').textContent=money(m.avg_income);result.innerHTML=data.report_html;debug.textContent=`征信文件：${m.credit_count||0}\n流水文件：${m.flow_file_count||0}\n唯一流水：${m.transaction_count||0} 笔\n跨文件重复：${m.duplicate_count||0} 笔\n处理异常：${(data.errors||[]).length} 个`}
async function download(format){if(!currentJobId||!lastData)return;statusEl.textContent='正在生成下载文件...';const res=await fetch('/export',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({job_id:currentJobId,format})});if(!res.ok){const e=await res.json();statusEl.textContent=e.error||'下载失败';return}const blob=await res.blob();const match=/filename\*=UTF-8''([^;]+)/i.exec(res.headers.get('Content-Disposition')||'');const name=match?decodeURIComponent(match[1]):`综合分析报告.${format}`;const url=URL.createObjectURL(blob),a=document.createElement('a');a.href=url;a.download=name;document.body.appendChild(a);a.click();a.remove();URL.revokeObjectURL(url);statusEl.textContent=`已生成：${name}`}
document.getElementById('pdfBtn').onclick=()=>download('pdf');document.getElementById('wordBtn').onclick=()=>download('word');document.getElementById('excelBtn').onclick=()=>download('excel');document.getElementById('printBtn').onclick=()=>window.print();document.getElementById('resetBtn').onclick=()=>{document.getElementById('form').reset();currentJobId=null;lastData=null;setDownloads(false);setProgress(0);updateFiles();statusEl.textContent='请选择客户文件夹。';result.innerHTML='<div class="empty">选择客户文件夹后，这里会显示征信、流水与综合风险提示。</div>';debug.textContent='等待分析。'};setDownloads(false);
</script></body></html>"""


class Handler(BaseHTTPRequestHandler):
    def send_json(self, data: dict, status: int = 200) -> None:
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/progress":
            job_id = (parse_qs(parsed.query).get("job_id") or [""])[0]
            job = get_job(job_id)
            self.send_json(job or {"error": "任务不存在"}, 200 if job else 404)
            return
        if parsed.path in {"/", "/comprehensive"}:
            body = INDEX.encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return
        self.send_error(404)

    def do_POST(self) -> None:
        if self.path == "/export":
            try:
                length = int(self.headers.get("Content-Length") or "0")
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                job = get_job(normalize_text(payload.get("job_id")))
                if job.get("status") != "done" or not job.get("result"):
                    raise ValueError("综合分析任务尚未完成。")
                content, filename, content_type = build_export(job["result"], payload.get("format") or "")
                self.send_response(200)
                self.send_header("Content-Type", content_type)
                self.send_header("Content-Length", str(len(content)))
                self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{quote(filename)}")
                self.end_headers()
                self.wfile.write(content)
            except Exception as exc:
                self.send_json({"error": str(exc), "trace": traceback.format_exc()}, 500)
            return
        if self.path != "/analyze/start":
            self.send_error(404)
            return
        try:
            form = cgi.FieldStorage(fp=self.rfile, headers=self.headers, environ={"REQUEST_METHOD": "POST"})
            if "file" not in form:
                raise ValueError("未收到文件夹内容。")
            items = form["file"] if isinstance(form["file"], list) else [form["file"]]
            job_id = uuid.uuid4().hex
            job_dir = UPLOADS / job_id
            job_dir.mkdir(parents=True, exist_ok=True)
            manifest = []
            used_names = set()
            for idx, item in enumerate(items, start=1):
                if not getattr(item, "filename", None):
                    continue
                relative = str(item.filename).replace("\\", "/")
                original_name = PurePosixPath(relative).name or f"upload-{idx}"
                saved_name = original_name
                counter = 1
                while saved_name in used_names:
                    saved_name = f"{Path(original_name).stem}-{counter}{Path(original_name).suffix}"
                    counter += 1
                used_names.add(saved_name)
                path = job_dir / saved_name
                path.write_bytes(item.file.read())
                manifest.append({"path": str(path), "name": original_name, "relative_path": relative, "status": "等待分类"})
            if not manifest:
                raise ValueError("文件夹中没有可读取的文件。")
            folder_name = normalize_text(form.getfirst("folder_name", "")) or "客户文件夹"
            passwords = flow.parse_passwords(form.getfirst("passwords", ""))
            set_job(job_id, status="queued", percent=0, message=f"已接收 {len(manifest)} 个文件")
            thread = threading.Thread(target=run_comprehensive_job, args=(job_id, manifest, folder_name, passwords), daemon=True)
            thread.start()
            self.send_json({"job_id": job_id, "status": "queued"})
        except Exception as exc:
            self.send_json({"error": str(exc), "trace": traceback.format_exc()}, 500)

    def log_message(self, format, *args):
        return


if __name__ == "__main__":
    port = int(__import__("os").environ.get("PORT", "8792"))
    server = ThreadingHTTPServer(("127.0.0.1", port), Handler)
    print(f"本地综合分析平台：http://127.0.0.1:{port}/comprehensive", flush=True)
    server.serve_forever()
